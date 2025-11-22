# AI-Powered Fund Compliance Agent
# Checks: Registration + Disclaimers + Structure + Securities/Values
# Local version with environment variable support

import pandas as pd
import re
import json
import os
import io
import time
from collections import Counter
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

print("="*70)
print("AI-POWERED FUND COMPLIANCE AGENT")
print("Checking: Registration + Disclaimers + Structure + Securities/Values")
print("="*70)

# ============================================================================
# STEP 1: Configure Gemini API
# ============================================================================
try:
    import google.generativeai as genai
    
    GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY not found in environment variables")
    
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel('gemini-2.5-flash')
    print("\nâœ“ Gemini API configured successfully")
except Exception as e:
    print(f"\nâš ï¸  Warning: Could not configure Gemini API: {e}")
    print("   Natural language queries will not be available")
    model = None

# ============================================================================
# STEP 2: Load metadata.json
# ============================================================================
metadata = None
try:
    if os.path.exists('metadata.json'):
        with open('metadata.json', 'r', encoding='utf-8') as f:
            metadata = json.load(f)
        print("âœ“ Loaded metadata.json")
        print(f"  - Management Company: {metadata.get('SociÃ©tÃ© de Gestion', 'N/A')}")
        print(f"  - Client Type: {'Professional' if metadata.get('Le client est-il un professionnel') else 'Retail'}")
        print(f"  - New Product: {'Yes' if metadata.get('Le document fait-il rÃ©fÃ©rence Ã  un nouveau Produit') else 'No'}")
except Exception as e:
    print(f"âš ï¸  metadata.json not found")

# ============================================================================
# STEP 3: Load Structure Rules
# ============================================================================
structure_rules = []
try:
    if os.path.exists('structure_rules.json'):
        with open('structure_rules.json', 'r', encoding='utf-8') as f:
            structure_data = json.load(f)
            structure_rules = structure_data.get('rules', [])
        print(f"âœ“ Loaded structure_rules.json ({len(structure_rules)} rules)")
    else:
        print("âš ï¸  structure_rules.json not found")
except Exception as e:
    print(f"âš ï¸  Could not load structure rules: {e}")

# ============================================================================
# STEP 4: Load Values/Securities Rules
# ============================================================================
values_rules = []
try:
    if os.path.exists('values_rules.json'):
        with open('values_rules.json', 'r', encoding='utf-8') as f:
            values_data = json.load(f)
            values_rules = values_data.get('rules', [])
        print(f"âœ“ Loaded values_rules.json ({len(values_rules)} rules)")
    else:
        print("âš ï¸  values_rules.json not found")
except Exception as e:
    print(f"âš ï¸  Could not load values rules: {e}")

# ============================================================================
# STEP 5: Load ESG Rules
# ============================================================================
esg_rules = []
try:
    if os.path.exists('esg_rules.json'):
        with open('esg_rules.json', 'r', encoding='utf-8') as f:
            esg_data = json.load(f)
            # Skip ESG_001 (requires Cartographie)
            esg_rules = [r for r in esg_data.get('rules', []) if r.get('rule_id') != 'ESG_001']
        print(f"âœ“ Loaded esg_rules.json ({len(esg_rules)} rules, skipping ESG_001)")
    else:
        print("âš ï¸  esg_rules.json not found")
except Exception as e:
    print(f"âš ï¸  Could not load ESG rules: {e}")

# ============================================================================
# STEP 6: Load Registration CSV file
# ============================================================================
df_registration = None
try:
    if os.path.exists('registration.csv'):
        df_registration = pd.read_csv('registration.csv')
        print("âœ“ Loaded registration.csv")
        print(f"  - Total rows: {len(df_registration)}")
        print(f"  - Columns: {list(df_registration.columns)}")
except Exception as e:
    print(f"âš ï¸  Registration file not found or error loading: {e}")

# ============================================================================
# STEP 7: Load Performance Rules
# ============================================================================
performance_rules = []
try:
    if os.path.exists('performance_rules.json'):
        with open('performance_rules.json', 'r', encoding='utf-8') as f:
            perf_data = json.load(f)
            performance_rules = perf_data.get('rules', [])
        print(f"âœ“ Loaded performance_rules.json ({len(performance_rules)} rules)")
    else:
        print("âš ï¸  performance_rules.json not found")
except Exception as e:
    print(f"âš ï¸  Could not load performance rules: {e}")

# ============================================================================
# STEP 8: Load General Rules
# ============================================================================
general_rules = []
try:
    if os.path.exists('general_rules.json'):
        with open('general_rules.json', 'r', encoding='utf-8') as f:
            gen_data = json.load(f)
            general_rules = gen_data.get('rules', [])
        print(f"âœ“ Loaded general_rules.json ({len(general_rules)} rules)")
    else:
        print("âš ï¸  general_rules.json not found")
except Exception as e:
    print(f"âš ï¸  Could not load general rules: {e}")

# ============================================================================
# STEP 9: Load Disclaimers Glossary
# ============================================================================
disclaimers_db = {}
df_disclaimers = None

try:
    if os.path.exists('GLOSSAIRE DISCLAIMERS 20231122.xlsx'):
        df_disclaimers = pd.read_excel('GLOSSAIRE DISCLAIMERS 20231122.xlsx', sheet_name='ENGLISH')
        print("âœ“ Loaded GLOSSAIRE DISCLAIMERS 20231122.xlsx (English)")
    else:
        print("\nâš ï¸  GLOSSAIRE DISCLAIMERS 20231122.xlsx not found")
except Exception as e:
    print(f"âš ï¸  Could not load disclaimers: {e}")

# ============================================================================
# STEP 10: Configure Token Factory API (Optional)
# ============================================================================
TOKENFACTORY_API_KEY = None
tokenfactory_client = None

try:
    TOKENFACTORY_API_KEY = os.getenv('TOKENFACTORY_API_KEY')
    if TOKENFACTORY_API_KEY:
        import httpx
        from openai import OpenAI

        http_client = httpx.Client(verify=False)
        tokenfactory_client = OpenAI(
            api_key=TOKENFACTORY_API_KEY,
            base_url="https://tokenfactory.esprit.tn/api",
            http_client=http_client
        )
        print("âœ“ Token Factory API configured successfully")
        print("  Using model: hosted_vllm/Llama-3.1-70B-Instruct")
except Exception as e:
    print(f"âš ï¸  Token Factory API not configured: {e}")
    print("   Will use Gemini for prospectus parsing if available")

# ============================================================================
# STEP 11: Parse Disclaimers into Database
# ============================================================================
if df_disclaimers is not None:
    try:
        # Find the header row
        header_row = None
        for idx, row in df_disclaimers.iterrows():
            row_str = ' '.join([str(cell) for cell in row if pd.notna(cell)]).lower()
            if 'document type' in row_str or 'non professional' in row_str:
                header_row = idx
                break

        if header_row is not None:
            cols = df_disclaimers.columns
            doc_type_col = cols[0]  # Column A
            retail_col = cols[1]     # Column B
            prof_col = cols[2]       # Column C

            # Parse disclaimers starting after header
            for idx in range(header_row + 1, len(df_disclaimers)):
                row = df_disclaimers.iloc[idx]
                doc_type = str(row[doc_type_col]).strip() if pd.notna(row[doc_type_col]) else ""

                if not doc_type or doc_type == 'nan' or len(doc_type) < 3:
                    continue

                retail_text = str(row[retail_col]) if pd.notna(row[retail_col]) else ""
                prof_text = str(row[prof_col]) if pd.notna(row[prof_col]) else ""

                if len(retail_text) > 50 or len(prof_text) > 50:
                    disclaimers_db[doc_type] = {
                        'retail': retail_text if len(retail_text) > 50 else None,
                        'professional': prof_text if len(prof_text) > 50 else None
                    }

        print(f"âœ“ Parsed {len(disclaimers_db)} disclaimer types")
    except Exception as e:
        print(f"âš ï¸  Error parsing disclaimers: {e}")

# ============================================================================
# STEP 12: Load and Parse Prospectus
# ============================================================================
prospectus_data = None
prospectus_raw_text = None

# Define the expected fields for the prospectus data
EXPECTED_FIELDS = {
    "fund_name": "The exact, full name of the sub-fund.",
    "investment_strategy": "A brief, one or two-sentence summary of the investment strategy.",
    "asset_allocation": {
        "equities": "Range (e.g., 'X-Y%') or target ('e.g., 'At least X%') for equity allocation.",
        "bonds": "Range (e.g., 'X-Y%') or target for bond allocation. Use null if not specified."
    },
    "geographic_allocation": {
        "europe": "Range (e.g., 'X-Y%') or target for European allocation. Use null if not specified.",
        "us": "Range (e.g., 'X-Y%') or target for US allocation. Use null if not specified."
    },
    "minimum_investment": "The minimum investment amount, including currency (e.g., '1,000,000 USD'). Use null if not specified for retail investors.",
    "risk_profile": "A brief summary of the primary risks (e.g., 'Risk of capital loss, equity market risk').",
    "sri": "The Summary Risk Indicator (SRI) score, formatted as 'X/7'.",
    "benchmark": "The full, exact name of the benchmark index. Use null if there is no benchmark.",
    "benchmark_specification": "Details about the benchmark, e.g., 'dividends reinvested', 'Net Total Return'. Use null if no benchmark.",
    "portfolio_holdings_range": "The typical number of holdings or lines, e.g., '60-80 lines'. Use null if not specified.",
    "management_fees": "The management fee as a percentage, e.g., '0.75%'.",
    "performance_fees": "The performance fee as a percentage, e.g., '10%'. Use null if not specified.",
    "investment_objective": "The main investment objective of the fund."
}

def estimate_tokens(text):
    """Rough estimate of tokens (1 token â‰ˆ 4 characters for English text)."""
    return len(text) // 4

def split_document_intelligently(text, max_chars=80000):
    """Split document into chunks, trying to keep sections together."""
    if len(text) <= max_chars:
        return [text]

    chunks = []
    current_chunk = ""
    paragraphs = text.split('\n\n')

    for para in paragraphs:
        if len(current_chunk) + len(para) > max_chars and current_chunk:
            chunks.append(current_chunk)
            current_chunk = para
        else:
            current_chunk += "\n\n" + para if current_chunk else para

    if current_chunk:
        chunks.append(current_chunk)

    return chunks

def clean_json_response(text):
    """Clean the response text to extract valid JSON."""
    text = text.strip()

    # Remove markdown code blocks if present
    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]

    if text.endswith("```"):
        text = text[:-3]

    text = text.strip()

    # Try to find JSON object boundaries
    start_idx = text.find('{')
    end_idx = text.rfind('}')

    if start_idx != -1 and end_idx != -1:
        text = text[start_idx:end_idx + 1]

    return text

def parse_chunk_with_ai(chunk, chunk_num, total_chunks, use_tokenfactory=False):
    """Parse a single chunk of the document using either Token Factory or Gemini."""
    schema_description = json.dumps(EXPECTED_FIELDS, indent=2)

    if total_chunks == 1:
        context = "This is the complete prospectus document."
    else:
        context = f"This is part {chunk_num + 1} of {total_chunks} of the prospectus document."

    prompt = f"""You are a professional financial analyst. {context}

Extract all available information from this text and format it as a valid JSON object with these fields:

{schema_description}

CRITICAL INSTRUCTIONS:
- Return ONLY a valid JSON object, no markdown formatting, no code blocks, no explanations
- Start directly with {{ and end with }}
- If a specific piece of information is not found in THIS chunk, use null for that field
- Pay close attention to details like percentages and currency symbols
- Ensure all string values are properly quoted

--- DOCUMENT TEXT ---
{chunk}
--- END DOCUMENT TEXT ---

Provide ONLY the JSON output:"""

    print(f"  Processing chunk {chunk_num + 1}/{total_chunks}...", end=" ")

    try:
        if use_tokenfactory and tokenfactory_client:
            # Use Token Factory (Llama-3.1-70B)
            print("[Token Factory]")
            response = tokenfactory_client.chat.completions.create(
                model="hosted_vllm/Llama-3.1-70B-Instruct",
                messages=[
                    {"role": "system", "content": "You are a financial document parser. You extract structured data and return only valid JSON with no additional text."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=2000,
                top_p=0.9
            )
            return response.choices[0].message.content

        elif model:
            # Use Gemini
            print("[Gemini]")
            response = model.generate_content(prompt)
            if response and hasattr(response, 'text'):
                return response.text

        return None

    except Exception as e:
        print(f"    âš ï¸  Error: {e}")
        return None

def merge_parsed_results(results):
    """Merge multiple parsed JSON objects, preferring non-null values."""
    merged = {}

    for result in results:
        for key, value in result.items():
            if key not in merged or merged[key] is None:
                merged[key] = value
            elif isinstance(value, dict) and isinstance(merged[key], dict):
                for sub_key, sub_value in value.items():
                    if sub_key not in merged[key] or merged[key][sub_key] is None:
                        merged[key][sub_key] = sub_value

    return merged

try:
    if os.path.exists('prospectus.docx'):
        from docx import Document
        doc_prospectus = Document('prospectus.docx')

        # Extract text sections
        prospectus_raw_text = '\n'.join([para.text for para in doc_prospectus.paragraphs if para.text.strip()])

        # Also extract text from tables
        for table in doc_prospectus.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        prospectus_raw_text += '\n' + cell.text

        char_count = len(prospectus_raw_text)
        estimated_tokens = estimate_tokens(prospectus_raw_text)

        print(f"âœ“ Loaded prospectus.docx ({char_count:,} characters, ~{estimated_tokens:,} tokens)")

        # Try to parse with AI (Token Factory or Gemini)
        if (tokenfactory_client or model) and len(prospectus_raw_text) > 100:
            # Determine which API to use
            use_tokenfactory = tokenfactory_client is not None
            api_name = "Token Factory (Llama-3.1-70B)" if use_tokenfactory else "Gemini"

            print(f"ðŸ¤– Parsing prospectus with {api_name}...")

            try:
                MAX_CHARS = 80000  # Conservative limit

                if char_count > MAX_CHARS:
                    print(f"   Document is large. Splitting into chunks...")
                    chunks = split_document_intelligently(prospectus_raw_text, MAX_CHARS)
                    print(f"   Split into {len(chunks)} chunks.")

                    # Process each chunk
                    parsed_results = []
                    for i, chunk in enumerate(chunks):
                        chunk_tokens = estimate_tokens(chunk)
                        print(f"\n   Chunk {i+1}: ~{chunk_tokens:,} tokens")

                        result_text = parse_chunk_with_ai(chunk, i, len(chunks), use_tokenfactory)
                        if result_text:
                            try:
                                cleaned = clean_json_response(result_text)
                                parsed = json.loads(cleaned)
                                parsed_results.append(parsed)
                                print(f"      âœ“ Chunk {i+1} parsed successfully")
                            except json.JSONDecodeError as e:
                                print(f"      âš ï¸  JSON parsing error for chunk {i+1}: {e}")
                                continue

                        time.sleep(1)  # Brief pause between chunks

                    if parsed_results:
                        print("\n   Merging results from all chunks...")
                        prospectus_data = merge_parsed_results(parsed_results)
                        prospectus_data['raw_text'] = prospectus_raw_text

                        extracted_fields = sum(1 for k, v in prospectus_data.items()
                                             if k != 'raw_text' and v is not None and v != {} and v != "")

                        print(f"\nâœ“ Parsed prospectus with AI ({extracted_fields}/{len(EXPECTED_FIELDS)} fields extracted)")
                    else:
                        print("\nâš ï¸  Failed to parse any chunks - using raw text only")
                        prospectus_data = {'raw_text': prospectus_raw_text}

                else:
                    # Document is small enough to process in one go
                    print("   Document size is acceptable. Processing in single request...")
                    result_text = parse_chunk_with_ai(prospectus_raw_text, 0, 1, use_tokenfactory)

                    if result_text:
                        try:
                            cleaned = clean_json_response(result_text)
                            prospectus_data = json.loads(cleaned)
                            prospectus_data['raw_text'] = prospectus_raw_text

                            extracted_fields = sum(1 for k, v in prospectus_data.items()
                                                 if k != 'raw_text' and v is not None and v != {} and v != "")

                            print(f"\nâœ“ Parsed prospectus with AI ({extracted_fields}/{len(EXPECTED_FIELDS)} fields extracted)")
                        except json.JSONDecodeError as e:
                            print(f"\nâš ï¸  JSON parsing error: {e}")
                            prospectus_data = {'raw_text': prospectus_raw_text}
                    else:
                        print("\nâš ï¸  No response from AI - using raw text only")
                        prospectus_data = {'raw_text': prospectus_raw_text}

                # Display extracted information
                if prospectus_data and prospectus_data != {'raw_text': prospectus_raw_text}:
                    if prospectus_data.get('fund_name'):
                        print(f"  - Fund: {prospectus_data['fund_name']}")
                    if prospectus_data.get('sri'):
                        print(f"  - SRI: {prospectus_data['sri']}")
                    if prospectus_data.get('benchmark'):
                        benchmark = prospectus_data['benchmark']
                        print(f"  - Benchmark: {benchmark[:80]}{'...' if len(benchmark) > 80 else ''}")
                    if prospectus_data.get('management_fees'):
                        print(f"  - Management Fees: {prospectus_data['management_fees']}")
                    if prospectus_data.get('investment_objective'):
                        obj = prospectus_data['investment_objective']
                        print(f"  - Objective: {obj[:80]}{'...' if len(obj) > 80 else ''}")

            except KeyboardInterrupt:
                print(f"\n\nâš ï¸  Parsing interrupted by user")
                prospectus_data = {'raw_text': prospectus_raw_text}
                print("  â†’ Using raw text only")

            except Exception as e:
                print(f"\nâš ï¸  AI parsing failed: {type(e).__name__}: {str(e)}")
                prospectus_data = {'raw_text': prospectus_raw_text}
                print("  â†’ Falling back to raw text only")

        else:
            prospectus_data = {'raw_text': prospectus_raw_text}
            if not tokenfactory_client and not model:
                print("âš ï¸  Using raw text only (No AI API available)")
            else:
                print("âš ï¸  Using raw text only (insufficient content)")

    else:
        print("âš ï¸  prospectus.docx not found")

except Exception as e:
    print(f"\nâŒ Could not load prospectus: {type(e).__name__}: {str(e)}")
    import traceback
    traceback.print_exc()

# ============================================================================
# STEP 13: Initialize remaining variables
# ============================================================================
prospectus_rules = []
funds_db = []
all_countries = set()

# Load prospectus rules
try:
    if os.path.exists('prospectus_rules.json'):
        with open('prospectus_rules.json', 'r', encoding='utf-8') as f:
            prosp_data = json.load(f)
            prospectus_rules = prosp_data.get('rules', [])
        print(f"âœ“ Loaded prospectus_rules.json ({len(prospectus_rules)} rules)")
    else:
        print("âš ï¸  prospectus_rules.json not found")
except Exception as e:
    print(f"âš ï¸  Could not load prospectus rules: {e}")

# Parse registration data
if df_registration is not None:
    def parse_registration_csv(df):
        """Parse CSV to extract fund registration data"""
        funds_db = []
        all_countries = set()

        for idx, row in df.iterrows():
            try:
                isin = str(row['isin']).strip() if pd.notna(row['isin']) else None

                # Validate ISIN format
                if isin and re.match(r'^[A-Z]{2}[A-Z0-9]{10}$', isin):
                    # Parse authorized countries from the list
                    countries_str = str(row['authorized_countries_list']) if pd.notna(row['authorized_countries_list']) else ""
                    authorized_countries = [c.strip() for c in countries_str.split(',') if c.strip()]

                    # Add to all countries set
                    all_countries.update(authorized_countries)

                    if authorized_countries:
                        funds_db.append({
                            'fund_family': str(row['fund_family']).strip() if pd.notna(row['fund_family']) else "",
                            'fund_name': str(row['fund_name']).strip() if pd.notna(row['fund_name']) else "",
                            'share_class': str(row['share_class']).strip() if pd.notna(row['share_class']) else "",
                            'isin': isin,
                            'num_countries': int(row['num_countries']) if pd.notna(row['num_countries']) else len(authorized_countries),
                            'authorized_countries': authorized_countries
                        })
            except Exception as e:
                if idx < 5:  # Only show first few errors
                    print(f"âš ï¸  Row {idx} - Error: {e}")
                continue

        return funds_db, sorted(list(all_countries))

    funds_db, all_countries = parse_registration_csv(df_registration)
    print(f"âœ“ Parsed {len(funds_db)} fund share classes")
    print(f"âœ“ Found {len(all_countries)} unique countries")

# ============================================================================
# HELPER: Token Factory API Call with Debug Logging
# ============================================================================
def call_tokenfactory_with_debug(prompt, system_message, function_name="", show_prompt=True, max_tokens=1000):
    """
    Call Token Factory API with debug logging and robust JSON parsing
    
    Args:
        prompt: User prompt
        system_message: System message
        function_name: Name of calling function (for logging)
        show_prompt: Whether to print the prompt (False for prospectus parsing)
        max_tokens: Maximum tokens for response
    
    Returns:
        Parsed JSON dict or None
    """
    if not tokenfactory_client:
        return None
    
    if show_prompt and function_name:
        print(f"\n{'='*70}")
        print(f"ðŸ¤– TOKEN FACTORY CALL: {function_name}")
        print(f"{'='*70}")
        print(f"SYSTEM: {system_message[:100]}...")
        print(f"\nPROMPT (first 500 chars):")
        print(prompt)

        print(f"{'='*70}\n")
    
    try:
        response = tokenfactory_client.chat.completions.create(
            model="hosted_vllm/Llama-3.1-70B-Instruct",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=max_tokens
        )
        
        result_text = response.choices[0].message.content.strip()
        
        if show_prompt and function_name:
            print(f"ðŸ“¥ RESPONSE:")
            print(result_text)
            print()
        
        # Clean and parse JSON
        result_text = result_text.replace('```json', '').replace('```', '').strip()
        
        # Remove any text before first { and after last }
        start_idx = result_text.find('{')
        end_idx = result_text.rfind('}')
        
        if start_idx != -1 and end_idx != -1:
            json_str = result_text[start_idx:end_idx + 1]
            try:
                return json.loads(json_str)
            except json.JSONDecodeError as e:
                print(f"âš ï¸  JSON parsing error in {function_name}: {e}")
                print(f"   Attempted to parse: {json_str[:200]}...")
                # Try to fix common issues
                try:
                    # Fix trailing commas
                    json_str = re.sub(r',(\s*[}\]])', r'\1', json_str)
                    # Fix unescaped quotes in strings
                    json_str = re.sub(r'(?<!\\)"(?=\w)', r'\\"', json_str)
                    return json.loads(json_str)
                except Exception as fix_error:
                    print(f"   Failed to fix JSON: {fix_error}")
                    return None
        
        print(f"âš ï¸  No valid JSON found in response for {function_name}")
        return None
        
    except Exception as e:
        print(f"âš ï¸  Token Factory API error in {function_name}: {e}")
        return None

print("\n" + "="*70)
print("INITIALIZATION COMPLETE")
print("="*70)
print("\nTo use this agent:")
print("1. Ensure all required JSON/CSV files are in the same directory")
print("2. Set API keys in .env file:")
print("   GEMINI_API_KEY=your_key_here")
print("   TOKENFACTORY_API_KEY=your_key_here (optional)")
print("3. Import and use the checking functions")
print("="*70)

# ============================================================================
# ALL CHECKING FUNCTIONS FROM AGENT.PY
# ============================================================================


def find_country_in_slides(doc, country):
    """Find which slide mentions a specific country"""
    country_lower = country.lower()
    slides_found = []

    # Check page_de_garde
    if 'page_de_garde' in doc:
        text = json.dumps(doc['page_de_garde']).lower()
        if country_lower in text:
            slides_found.append(doc['page_de_garde'].get('slide_number', 1))

    # Check slide_2
    if 'slide_2' in doc:
        text = json.dumps(doc['slide_2']).lower()
        if country_lower in text:
            slides_found.append(doc['slide_2'].get('slide_number', 2))

    # Check pages_suivantes
    if 'pages_suivantes' in doc:
        for slide in doc['pages_suivantes']:
            text = json.dumps(slide).lower()
            if country_lower in text:
                slides_found.append(slide.get('slide_number', '?'))

    # Check page_de_fin
    if 'page_de_fin' in doc:
        text = json.dumps(doc['page_de_fin']).lower()
        if country_lower in text:
            slides_found.append(doc['page_de_fin'].get('slide_number', '?'))

    if slides_found:
        # Remove duplicates and sort
        slides_found = sorted(set([s for s in slides_found if s != '?']))
        if len(slides_found) == 1:
            return f"Slide {slides_found[0]}"
        else:
            return f"Slides {', '.join(map(str, slides_found))}"

    return "Multiple slides"

def extract_all_text_from_doc(doc):
    """Extract all text content from JSON document"""
    def extract_text(obj):
        if isinstance(obj, dict):
            texts = []
            for value in obj.values():
                texts.extend(extract_text(value))
            return texts
        elif isinstance(obj, list):
            texts = []
            for item in obj:
                texts.extend(extract_text(item))
            return texts
        elif isinstance(obj, str):
            return [obj]
        else:
            return []

    all_texts = extract_text(doc)
    return ' '.join(all_texts)

def extract_section_text(doc, section):
    """Extract text from a specific section (page_de_garde, slide_2, page_de_fin)"""
    if section in doc:
        return json.dumps(doc[section])
    return ""

def check_text_presence(text, keywords, case_sensitive=False):
    """Check if keywords are present in text"""
    if not case_sensitive:
        text = text.lower()
        keywords = [k.lower() for k in keywords]

    for keyword in keywords:
        if keyword in text:
            return True
    return False

def find_phrase_in_document(doc, phrase):
    """Find which slides contain a specific phrase"""
    phrase_lower = phrase.lower()
    locations = []

    sections = [
        ('page_de_garde', doc.get('page_de_garde', {})),
        ('slide_2', doc.get('slide_2', {})),
        ('page_de_fin', doc.get('page_de_fin', {}))
    ]

    for section_name, section_data in sections:
        if section_data:
            text = json.dumps(section_data).lower()
            if phrase_lower in text:
                slide_num = section_data.get('slide_number', '?')
                locations.append(f"Slide {slide_num}")

    # Check pages_suivantes
    if 'pages_suivantes' in doc:
        for slide in doc['pages_suivantes']:
            text = json.dumps(slide).lower()
            if phrase_lower in text:
                slide_num = slide.get('slide_number', '?')
                locations.append(f"Slide {slide_num}")

    if locations:
        return ', '.join(locations)
    return "Document-wide"

def check_disclaimer_in_document(document_text, required_disclaimer):
    """Use Gemini to check if a disclaimer is present in document (fuzzy match)"""
    if not model:
        return {
            'status': 'ERROR',
            'message': 'Gemini API not configured'
        }

    prompt = f"""You are a compliance checker. Check if a required disclaimer is present in a marketing document.

Required disclaimer:
---
{required_disclaimer[:1500]}
---

Document text (excerpt):
---
{document_text[:2000]}
---

Instructions:
1. Check if the MEANING and KEY POINTS of the required disclaimer are present
2. Allow for slight variations in wording
3. Key compliance points must be covered

Respond in JSON format:
{{
  "status": "FOUND" or "MISSING" or "PARTIAL",
  "confidence": 0-100,
  "explanation": "Brief explanation of your finding"
}}
"""

    try:
        response = model.generate_content(prompt)
        result = json.loads(response.text.strip().replace('```json', '').replace('```', ''))
        return result
    except Exception as e:
        return {
            'status': 'ERROR',
            'message': f'Error: {e}'
        }

def check_date_format(text):
    """Check if text contains a valid date in MM/YYYY or Month YYYY format"""
    patterns = [
        r'\d{2}/\d{4}',  # MM/YYYY
        r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}',  # Month YYYY
        r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}'  # Mon YYYY
    ]

    for pattern in patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True
    return False


# ============================================================================
# STEP XX: Prospectus Compliance Checker
# ============================================================================
def check_prospectus_compliance(doc, prospectus_data):
    """Check if marketing document matches prospectus data - COMPREHENSIVE"""
    violations = []

    if not prospectus_data:
        return violations

    doc_text = extract_all_text_from_doc(doc)
    doc_text_lower = doc_text.lower()

    # PROSP_001: Investment Strategy (GEN_008)
    if 'investment_strategy' in prospectus_data and prospectus_data['investment_strategy']:

        if model:
            prompt = f"""Compare investment strategy consistency:

Prospectus: {prospectus_data.get('investment_strategy', '')}

Document: {doc_text[:2500]}

Check:
1. Strategy description matches
2. Investment thresholds match
3. Asset allocation ranges match
4. Geographic allocation matches
5. Minimum investment matches

JSON response:
{{
  "consistent": true/false,
  "confidence": 0-100,
  "mismatches": ["list of specific mismatches"]
}}
"""
            try:
                response = model.generate_content(prompt)
                result = json.loads(response.text.strip().replace('```json', '').replace('```', ''))

                if not result.get('consistent', True):
                    violations.append({
                        'type': 'PROSPECTUS',
                        'severity': 'CRITICAL',
                        'slide': 'Strategy section',
                        'location': 'pages_suivantes',
                        'rule': 'PROSP_001: Investment strategy must match prospectus',
                        'message': 'Strategy presentation inconsistent with prospectus',
                        'evidence': '\n'.join(result.get('mismatches', ['Strategy mismatch'])),
                        'confidence': result.get('confidence', 0)
                    })
            except Exception as e:
                print(f"âš ï¸  Could not validate strategy: {e}")

    # PROSP_002 & PROSP_003: Risk Profile and SRI (STRUCT_010, GEN_004)
    if 'sri' in prospectus_data and prospectus_data['sri']:
        prospectus_sri = prospectus_data['sri'].lower()

        # Check SRI presence
        sri_pattern = r'sri[:\s]*(\d)/7'
        sri_matches = re.findall(sri_pattern, doc_text_lower)

        if not sri_matches:
            violations.append({
                'type': 'PROSPECTUS',
                'severity': 'CRITICAL',
                'slide': 'SRI section',
                'location': 'Missing',
                'rule': 'PROSP_003: SRI must be displayed with disclaimer',
                'message': f"SRI not found in document",
                'evidence': f"Prospectus SRI: {prospectus_data['sri']}. Must include SRI with disclaimer on same slide"
            })
        else:
            doc_sri = f"{sri_matches[0]}/7"
            if doc_sri != prospectus_sri:
                violations.append({
                    'type': 'PROSPECTUS',
                    'severity': 'CRITICAL',
                    'slide': 'SRI section',
                    'location': 'Multiple locations',
                    'rule': 'PROSP_003: SRI must match prospectus exactly',
                    'message': f"SRI mismatch: Document '{doc_sri}' â‰  Prospectus '{prospectus_sri}'",
                    'evidence': f"Update all SRI mentions to: {prospectus_data['sri']}"
                })

    if 'risk_profile' in prospectus_data and prospectus_data['risk_profile']:

        # Check if risk description is present
        risk_terms = ['risk', 'volatility', 'loss', 'capital']
        has_risk = any(term in doc_text_lower for term in risk_terms)

        if not has_risk:
            violations.append({
                'type': 'PROSPECTUS',
                'severity': 'CRITICAL',
                'slide': 'Slide 2 or Risk section',
                'location': 'Missing',
                'rule': 'PROSP_002: Exhaustive risk profile required',
                'message': 'Risk profile description appears incomplete or missing',
                'evidence': 'Must include exhaustive risk profile conforming to prospectus'
            })

    # PROSP_004, PROSP_005, PROSP_006: Benchmark (PERF_014-016, PERF_019)
    if 'benchmark' in prospectus_data and prospectus_data['benchmark']:
        prospectus_benchmark = prospectus_data['benchmark']

        # Check if document mentions performance
        has_performance = any(kw in doc_text_lower for kw in ['performance', 'return', 'yield'])

        if has_performance:
            # Check if prospectus benchmark is used
            if prospectus_benchmark.lower() not in doc_text_lower:
                violations.append({
                    'type': 'PROSPECTUS',
                    'severity': 'CRITICAL',
                    'slide': 'Performance section',
                    'location': 'pages_suivantes',
                    'rule': 'PROSP_004: Must use official prospectus benchmark',
                    'message': 'Performance shown without prospectus benchmark or with wrong benchmark',
                    'evidence': f"Required benchmark: {prospectus_benchmark}. ONLY this benchmark allowed (PERF_016)"
                })

            # Check benchmark specifications (dividends, etc.)
            if 'benchmark_specification' in prospectus_data:
                spec = prospectus_data['benchmark_specification'].lower()
                if spec not in doc_text_lower:
                    violations.append({
                        'type': 'PROSPECTUS',
                        'severity': 'MAJOR',
                        'slide': 'Performance section',
                        'location': 'pages_suivantes',
                        'rule': 'PROSP_005: Benchmark specifications must match prospectus',
                        'message': 'Benchmark specification missing or incorrect',
                        'evidence': f"Required: {prospectus_data['benchmark_specification']} (e.g., dividends reinvested)"
                    })

    # PROSP_006: Performance target
    if 'performance_target' in prospectus_data and prospectus_data['performance_target']:
        target = prospectus_data['performance_target']

        # If target exists in prospectus, it must be shown
        if target.lower() not in doc_text_lower:
            violations.append({
                'type': 'PROSPECTUS',
                'severity': 'CRITICAL',
                'slide': 'Performance section',
                'location': 'pages_suivantes',
                'rule': 'PROSP_006: Must compare to performance target if defined in prospectus',
                'message': 'Performance target from prospectus not mentioned',
                'evidence': f"Prospectus target: {target}. Must show comparison to this target (PERF_015)"
            })

    # PROSP_007: Portfolio holdings (STRUCT_023)
    holdings_terms = ['lines', 'holdings', 'positions', 'securities in portfolio', 'number of']
    mentions_holdings = any(term in doc_text_lower for term in holdings_terms)
    if mentions_holdings:
        if 'portfolio_holdings_range' in prospectus_data and prospectus_data['portfolio_holdings_range']:
            # Holdings range exists - this is fine
            pass
        elif 'portfolio_holdings_range' in prospectus_data:
            # Field exists but is None/empty
            violations.append({
                    'type': 'PROSPECTUS',
                    'severity': 'MAJOR',
                    'slide': 'Portfolio section',
                    'location': 'pages_suivantes',
                    'rule': 'PROSP_007: Portfolio holdings count must be in prospectus if mentioned',
                    'message': 'Document mentions portfolio holdings count but prospectus does not specify this',
                    'evidence': 'Remove holdings count OR ensure it is stated in prospectus (STRUCT_023)'
                })
        else:
            # Prospectus data doesn't have this field - flag warning
            violations.append({
                'type': 'PROSPECTUS',
                'severity': 'WARNING',
                'slide': 'Portfolio section',
                'location': 'pages_suivantes',
                'rule': 'PROSP_007: Verify portfolio holdings in prospectus',
                'message': 'Document mentions portfolio holdings - verify this is in prospectus',
                'evidence': 'Number of portfolio lines can only be mentioned if stated in prospectus'
            })

    # PROSP_009: Asset allocation ranges
    if 'asset_allocation' in prospectus_data and prospectus_data['asset_allocation']:
        allocation = prospectus_data['asset_allocation']

        if model:
            prompt = f"""Check asset allocation consistency:

Prospectus allocation: {json.dumps(allocation)}

Document text: {doc_text[:2000]}

Are the ranges consistent? JSON:
{{
  "consistent": true/false,
  "mismatches": ["specific mismatches"]
}}
"""
            try:
                response = model.generate_content(prompt)
                result = json.loads(response.text.strip().replace('```json', '').replace('```', ''))

                if not result.get('consistent', True):
                    violations.append({
                        'type': 'PROSPECTUS',
                        'severity': 'CRITICAL',
                        'slide': 'Strategy/allocation section',
                        'location': 'pages_suivantes',
                        'rule': 'PROSP_009: Asset allocation must match prospectus',
                        'message': 'Asset allocation ranges inconsistent with prospectus',
                        'evidence': '\n'.join(result.get('mismatches', ['Allocation mismatch']))
                    })
            except:
                pass

    # PROSP_010: Geographic allocation
    if 'geographic_allocation' in prospectus_data and prospectus_data['geographic_allocation']:

        geo_allocation = prospectus_data['geographic_allocation']

        if model:
            prompt = f"""Check geographic allocation consistency:

Prospectus: {json.dumps(geo_allocation)}

Document: {doc_text[:2000]}

Consistent? JSON: {{"consistent": true/false, "mismatches": []}}
"""
            try:
                response = model.generate_content(prompt)
                result = json.loads(response.text.strip().replace('```json', '').replace('```', ''))

                if not result.get('consistent', True):
                    violations.append({
                        'type': 'PROSPECTUS',
                        'severity': 'CRITICAL',
                        'slide': 'Geographic allocation section',
                        'location': 'pages_suivantes',
                        'rule': 'PROSP_010: Geographic allocation must match prospectus',
                        'message': 'Geographic allocation inconsistent with prospectus',
                        'evidence': '\n'.join(result.get('mismatches', []))
                    })
            except:
                pass

    # PROSP_011: Investment objective
    if 'investment_objective' in prospectus_data and prospectus_data['investment_objective']:

        objective = prospectus_data['investment_objective']

        # Simple keyword check
        if len(objective) > 20:
            key_words = objective.lower().split()[:10]  # First 10 words
            matches = sum(1 for word in key_words if len(word) > 4 and word in doc_text_lower)

            if matches < 5:  # Less than half match
                violations.append({
                    'type': 'PROSPECTUS',
                    'severity': 'CRITICAL',
                    'slide': 'Objective section',
                    'location': 'pages_suivantes',
                    'rule': 'PROSP_011: Investment objective must match prospectus',
                    'message': 'Investment objective wording differs significantly from prospectus',
                    'evidence': f"Prospectus objective: {objective[:200]}..."
                })

    # PROSP_012: Minimum investment
    if 'minimum_investment' in prospectus_data and prospectus_data['minimum_investment']:

        min_inv = str(prospectus_data['minimum_investment'])

        if 'minimum' in doc_text_lower and min_inv not in doc_text:
            violations.append({
                'type': 'PROSPECTUS',
                'severity': 'MAJOR',
                'slide': 'Fund characteristics',
                'location': 'pages_suivantes or page_de_fin',
                'rule': 'PROSP_012: Minimum investment must match prospectus',
                'message': 'Minimum investment amount differs from prospectus',
                'evidence': f"Prospectus minimum: {prospectus_data['minimum_investment']}"
            })

    # PROSP_013: Management fees
    if 'management_fees' in prospectus_data and prospectus_data['management_fees']:

        fees = str(prospectus_data['management_fees'])

        if 'fees' in doc_text_lower or 'charges' in doc_text_lower:
            if fees not in doc_text:
                violations.append({
                    'type': 'PROSPECTUS',
                    'severity': 'CRITICAL',
                    'slide': 'Fees section',
                    'location': 'pages_suivantes or page_de_fin',
                    'rule': 'PROSP_013: Management fees must match prospectus',
                    'message': 'Management fees differ from prospectus',
                    'evidence': f"Prospectus fees: {prospectus_data['management_fees']}"
                })

    # PROSP_014: Fund characteristics completeness
    # Check if end section has sufficient detail
    if 'page_de_fin' in doc:
        end_text = extract_section_text(doc, 'page_de_fin')

        required_terms = ['isin', 'share class', 'fees', 'nav', 'currency']
        missing_terms = [term for term in required_terms if term not in end_text.lower()]

        if len(missing_terms) >= 3:  # Missing most characteristics
            violations.append({
                'type': 'PROSPECTUS',
                'severity': 'MAJOR',
                'slide': 'End of presentation',
                'location': 'page_de_fin',
                'rule': 'PROSP_014: Detailed fund characteristics required at end',
                'message': 'Fund characteristics section appears incomplete',
                'evidence': f"Missing elements: {', '.join(missing_terms)}. Must include detailed characteristics"
            })

    # PROSP_008: COMPREHENSIVE DATA CHECK (STRUCT_024)
    # This is a catch-all - flag for manual review
    violations.append({
        'type': 'PROSPECTUS',
        'severity': 'WARNING',
        'slide': 'Document-wide',
        'location': 'All data points',
        'rule': 'PROSP_008: Verify ALL data consistency with legal docs',
        'message': 'âš ï¸ MANUAL REVIEW REQUIRED: Verify all data matches KID, Prospectus, SFDR Annex',
        'evidence': 'All numerical data, percentages, dates, and facts must be verified against legal documentation (STRUCT_024)'
    })

    return violations

# ============================================================================
# STEP 11: New LLM APPROACH
# ============================================================================
def check_country_context_with_llm(doc_text, country, mention_context, authorized_countries):
    """
    Use LLM to determine if country mention is a DISTRIBUTION claim

    Args:
        doc_text: Full document text
        country: Country name to check
        mention_context: 500 chars around the mention
        authorized_countries: List of authorized countries

    Returns:
        dict with is_distribution_claim, context, confidence
    """
    if not tokenfactory_client:
        return {'is_distribution_claim': True, 'confidence': 50}  # Fallback

    prompt = f"""You are a financial compliance expert. Analyze if this country mention represents a DISTRIBUTION/MARKETING claim.

COUNTRY: {country}

CONTEXT AROUND MENTION:
{mention_context}

AUTHORIZED COUNTRIES: {', '.join(authorized_countries[:10])}

QUESTION: Is "{country}" mentioned in a DISTRIBUTION/AUTHORIZATION context?

Consider these contexts:
- âœ… DISTRIBUTION: "authorized in", "distributed in", "available in", "marketed in", "registered in", "autorisÃ© en", "distribuÃ© en", "commercialisÃ© en"
- âŒ NOT DISTRIBUTION: "invests in [country] equities", "exposure to [country] markets", "risk related to [country]", "domiciled in [country]", "NOT authorized in", "non autorisÃ© en"

Also consider:
- Negations: "NOT authorized" or "non autorisÃ©" = NOT a distribution claim
- Investment universe: "invests in German stocks" = NOT a distribution claim
- Risk exposure: "US market risk" = NOT a distribution claim
- Legal domicile: "company established in France" = NOT a distribution claim

Respond ONLY with valid JSON:
{{
  "is_distribution_claim": true or false,
  "context_type": "authorized_distribution" or "investment_universe" or "risk_exposure" or "legal_domicile" or "negation" or "example",
  "evidence": "exact phrase from context",
  "confidence": 0-100,
  "reasoning": "brief explanation"
}}"""

    result = call_tokenfactory_with_debug(
        prompt=prompt,
        system_message="You are a financial compliance expert. Analyze context carefully. Respond ONLY with valid JSON, no markdown.",
        function_name=f"check_country_context ({country})",
        show_prompt=True
    )
    
    if result:
        return result
    else:
        print(f"    âš ï¸  LLM error for {country}: Failed to parse response")
        return {'is_distribution_claim': True, 'confidence': 50}  # Conservative fallback


def extract_mention_context(doc_text, country, context_chars=500):
    """Extract text around country mention for context analysis"""
    doc_lower = doc_text.lower()
    country_lower = country.lower()

    # Find all occurrences
    contexts = []
    start = 0
    while True:
        pos = doc_lower.find(country_lower, start)
        if pos == -1:
            break

        context_start = max(0, pos - context_chars)
        context_end = min(len(doc_text), pos + len(country) + context_chars)
        context = doc_text[context_start:context_end]
        contexts.append(context)

        start = pos + 1

    return contexts


def check_registration_rules_enhanced(doc, fund_isin, authorized_countries):
    """
    Enhanced registration check using LLM for context-aware country detection

    Returns:
        list of violations with LLM confidence scores
    """
    violations = []

    if not fund_isin or not authorized_countries:
        return violations

    doc_text = extract_all_text_from_doc(doc)

    # Detect country mentions using smart matching
    mentioned_countries = {}

    # Check each country in database
    for country in all_countries:
        # Get variations
        country_variations = [
            country.lower(),
            country.lower().replace(' (fund)', ''),
            country.lower().split()[0]
        ]

        found = False
        for variation in country_variations:
            if len(variation) > 3 and variation in doc_text.lower():
                found = True
                break

        if found:
            # Get contexts for LLM analysis
            contexts = extract_mention_context(doc_text, country)

            # Analyze each context
            for context in contexts:
                llm_result = check_country_context_with_llm(
                    doc_text,
                    country,
                    context,
                    authorized_countries
                )

                # Only flag if LLM confirms it's a distribution claim
                if llm_result.get('is_distribution_claim') and llm_result.get('confidence', 0) > 60:
                    mentioned_countries[country] = {
                        'context_type': llm_result.get('context_type'),
                        'evidence': llm_result.get('evidence'),
                        'confidence': llm_result.get('confidence'),
                        'reasoning': llm_result.get('reasoning')
                    }
                    break

    print(f"ðŸ” LLM-analyzed countries: {len(mentioned_countries)}")
    if mentioned_countries:
        print(f"   Distribution claims detected: {', '.join(mentioned_countries.keys())}\n")

    # Check authorization
    unauthorized_countries = {}

    for country, details in mentioned_countries.items():
        authorized = False
        for auth_country in authorized_countries:
            country_clean = country.lower().replace(' (fund)', '').strip()
            auth_clean = auth_country.lower().replace(' (fund)', '').strip()

            if country_clean == auth_clean or country_clean in auth_clean or auth_clean in country_clean:
                authorized = True
                break

        if not authorized:
            slide_location = find_country_in_slides(doc, country)
            unauthorized_countries[country] = {
                'location': slide_location,
                'confidence': details['confidence'],
                'evidence': details['evidence']
            }

    if unauthorized_countries:
        country_list = []
        for country, info in sorted(unauthorized_countries.items()):
            country_list.append(
                f"   â€¢ {country} ({info['location']}) [Confidence: {info['confidence']}%]\n"
                f"     Evidence: \"{info['evidence'][:100]}...\""
            )

        violations.append({
            'type': 'REGISTRATION',
            'severity': 'CRITICAL',
            'slide': 'Multiple slides',
            'location': 'Document-wide',
            'rule': 'Countries must match fund registration data',
            'message': f"Document mentions {len(unauthorized_countries)} unauthorized distribution {'claim' if len(unauthorized_countries) == 1 else 'claims'}:\n" + '\n'.join(country_list),
            'evidence': f"Fund {fund_isin} is only authorized in: {', '.join(sorted(authorized_countries)[:10])}{'...' if len(authorized_countries) > 10 else ''}"
        })

    return violations
def llm_check_structural_element(section_text, requirement_type, expected_value=None, metadata=None):
    """
    Universal LLM-based structural element validator

    Args:
        section_text: Text from document section
        requirement_type: 'fund_name', 'date', 'promotional_mention', 'target_audience',
                         'confidentiality', 'legal_mention', 'risk_profile'
        expected_value: Expected value (e.g., fund name from prospectus)
        metadata: Additional context (client_type, fund_status, etc.)

    Returns:
        dict with found, confidence, evidence, alternatives
    """
    if not tokenfactory_client:
        return {'found': False, 'confidence': 50, 'evidence': 'LLM not available'}

    # Build requirement-specific prompts
    prompts = {
        'fund_name': f"""Does this section contain the FULL fund name?

SECTION TEXT:
{section_text[:1500]}

EXPECTED FUND NAME: {expected_value or 'Unknown'}

Look for:
- Full fund name (e.g., "ODDO BHF Algo Trend US")
- Compartment/sub-fund name
- SICAV/FCP name
- Any variation in English or French

IGNORE:
- Just the word "fund" or "fonds" without actual name
- Company name only (e.g., just "ODDO BHF")

Respond JSON:
{{
  "found": true/false,
  "found_name": "exact text found or null",
  "is_complete_name": true/false,
  "confidence": 0-100,
  "evidence": "where/how found"
}}""",

        'date': f"""Does this section contain a valid date in acceptable format?

SECTION TEXT:
{section_text[:1000]}

ACCEPTABLE FORMATS:
- MM/YYYY or MM-YYYY (e.g., "01/2025", "01-2025")
- Month YYYY in English (e.g., "January 2025", "Jan 2025")
- Month YYYY in French (e.g., "Janvier 2025", "janv. 2025")
- YYYY/MM or YYYY-MM (e.g., "2025/01", "2025-01")
- Quarter format (e.g., "Q1 2025", "1er trimestre 2025")
- DD.MM.YYYY European format (e.g., "31.01.2025")

Respond JSON:
{{
  "found": true/false,
  "date_found": "exact date text",
  "format": "MM/YYYY or other",
  "confidence": 0-100
}}""",

        'promotional_mention': f"""Does this section indicate this is a PROMOTIONAL/MARKETING document?

SECTION TEXT:
{section_text[:1000]}

Look for expressions like:
- English: "promotional document", "marketing material", "commercial document", "advertising material"
- French: "document promotionnel", "document Ã  caractÃ¨re commercial", "matÃ©riel de marketing", "communication commerciale", "document publicitaire", "support de vente"

Respond JSON:
{{
  "found": true/false,
  "expression_found": "exact phrase",
  "language": "en/fr",
  "confidence": 0-100
}}""",

        'target_audience': f"""Does this section specify the TARGET AUDIENCE (retail vs professional investors)?

SECTION TEXT:
{section_text[:1000]}

Look for:
- English: "retail investors", "professional investors", "qualified investors", "institutional investors", "eligible counterparties"
- French: "investisseurs non-professionnels", "investisseurs professionnels", "clientÃ¨le de dÃ©tail", "investisseurs qualifiÃ©s", "clients qualifiÃ©s", "investisseurs avertis"

Respond JSON:
{{
  "found": true/false,
  "audience_type": "retail" or "professional" or "both" or null,
  "expression_found": "exact phrase",
  "confidence": 0-100
}}""",

        'confidentiality': f"""Does this section contain a CONFIDENTIALITY/NON-DISCLOSURE notice?

SECTION TEXT:
{section_text[:1000]}

Look for:
- English: "confidential", "do not disclose", "do not distribute", "strictly confidential", "proprietary information", "for private use only", "internal use only"
- French: "confidentiel", "ne pas diffuser", "ne pas divulguer", "strictement confidentiel", "usage interne uniquement", "document confidentiel"

Respond JSON:
{{
  "found": true/false,
  "expression_found": "exact phrase",
  "confidence": 0-100
}}""",

        'legal_mention': f"""Does this section contain the LEGAL MENTION of the management company?

SECTION TEXT:
{section_text[:1500]}

Look for mentions of:
- Management company name (e.g., "ODDO BHF Asset Management", "ODDO BHF AM")
- English: "management company", "asset management", "investment manager", "portfolio manager", "managed by"
- French: "sociÃ©tÃ© de gestion", "sociÃ©tÃ© de gestion de portefeuille", "SGP", "gÃ©rÃ© par", "gestionnaire"

Must include both: company name AND role/status

Respond JSON:
{{
  "found": true/false,
  "company_name": "name found",
  "expression_found": "full legal mention",
  "confidence": 0-100
}}""",

        'risk_profile': f"""Does this section contain adequate RISK PROFILE information?

SECTION TEXT:
{section_text[:1500]}

Look for risk-related content:
- Risk keywords: "risk", "risque", "volatility", "volatilitÃ©", "loss", "perte", "capital"
- Risk types: market risk, liquidity risk, currency risk, etc.
- Risk warnings: capital at risk, risk of loss, etc.

Respond JSON:
{{
  "found": true/false,
  "risk_elements_found": ["list of risk mentions"],
  "appears_comprehensive": true/false,
  "confidence": 0-100
}}"""
    }

    prompt = prompts.get(requirement_type)
    if not prompt:
        return {'found': False, 'confidence': 0, 'evidence': 'Unknown requirement type'}

    result = call_tokenfactory_with_debug(
        prompt=prompt,
        system_message="You are a multilingual financial compliance expert. Understand context and semantic meaning, not just exact keywords. Respond ONLY with valid JSON.",
        function_name=f"check_structural_element ({requirement_type})",
        show_prompt=True
    )
    
    if result:
        return result
    else:
        print(f"    âš ï¸  LLM error for {requirement_type}: Failed to parse response")
        return {'found': False, 'confidence': 50, 'evidence': 'JSON parsing failed'}


def check_structure_rules_enhanced(doc, client_type, fund_status='active', prospectus_data=None):
    """
    Enhanced structure rules checker using LLM for element detection

    Returns:
        list of violations with LLM-validated evidence
    """
    violations = []

    if not structure_rules:
        return violations

    section_mapping = {
        'Cover Page': 'page_de_garde',
        'Disclaimer Slide (Slide 2)': 'slide_2',
        'Back Page': 'page_de_fin'
    }

    for rule in structure_rules:
        # Check if rule applies
        applies = False
        rule_applies_to = rule.get('applies_to', {})

        if client_type.lower() in [ct.lower() for ct in rule_applies_to.get('client_type', [])]:
            if 'all' in rule_applies_to.get('fund_status', []) or fund_status in rule_applies_to.get('fund_status', []):
                applies = True

        if not applies:
            continue

        # Get section
        section = rule.get('section', '')
        json_key = section_mapping.get(section, '')
        section_text = extract_section_text(doc, json_key)

        rule_id = rule.get('rule_id', '')
        severity = rule.get('severity', 'major').upper()
        rule_text_lower = rule.get('rule_text', '').lower()

        # Use LLM validation based on rule type
        if 'fund name' in rule_text_lower:
            expected_name = prospectus_data.get('fund_name') if prospectus_data else None
            result = llm_check_structural_element(section_text, 'fund_name', expected_name)

            if not result.get('found') or result.get('confidence', 0) < 70:
                violations.append({
                    'type': 'STRUCTURE',
                    'severity': severity,
                    'slide': section,
                    'location': json_key,
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Fund name missing or incomplete (confidence: {result.get('confidence', 0)}%)",
                    'evidence': result.get('evidence', 'Fund name not clearly present')
                })

        elif 'month and year' in rule_text_lower or 'date' in rule_text_lower:
            result = llm_check_structural_element(section_text, 'date')

            if not result.get('found') or result.get('confidence', 0) < 70:
                violations.append({
                    'type': 'STRUCTURE',
                    'severity': severity,
                    'slide': section,
                    'location': json_key,
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Date missing or in wrong format (confidence: {result.get('confidence', 0)}%)",
                    'evidence': f"Required: MM/YYYY, Month YYYY, or equivalent. {result.get('evidence', '')}"
                })

        elif 'promotional document' in rule_text_lower:
            result = llm_check_structural_element(section_text, 'promotional_mention')

            if not result.get('found') or result.get('confidence', 0) < 70:
                violations.append({
                    'type': 'STRUCTURE',
                    'severity': severity,
                    'slide': section,
                    'location': json_key,
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Promotional document mention missing (confidence: {result.get('confidence', 0)}%)",
                    'evidence': rule.get('remediation', 'Must clearly state this is promotional/marketing material')
                })

        elif 'target audience' in rule_text_lower:
            result = llm_check_structural_element(section_text, 'target_audience')

            if not result.get('found') or result.get('confidence', 0) < 70:
                violations.append({
                    'type': 'STRUCTURE',
                    'severity': severity,
                    'slide': section,
                    'location': json_key,
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Target audience not specified (confidence: {result.get('confidence', 0)}%)",
                    'evidence': f"Must specify: retail/professional/qualified investors"
                })

        elif 'do not disclose' in rule_text_lower and client_type.lower() == 'professional':
            result = llm_check_structural_element(section_text, 'confidentiality')

            if not result.get('found') or result.get('confidence', 0) < 70:
                violations.append({
                    'type': 'STRUCTURE',
                    'severity': severity,
                    'slide': section,
                    'location': json_key,
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Confidentiality notice missing for professional document (confidence: {result.get('confidence', 0)}%)",
                    'evidence': rule.get('remediation', 'Must include confidentiality/non-disclosure notice')
                })

        elif 'legal mention' in rule_text_lower or 'management company' in rule_text_lower:
            result = llm_check_structural_element(section_text, 'legal_mention')

            if not result.get('found') or result.get('confidence', 0) < 70:
                violations.append({
                    'type': 'STRUCTURE',
                    'severity': severity,
                    'slide': section,
                    'location': json_key,
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Legal mention of management company missing (confidence: {result.get('confidence', 0)}%)",
                    'evidence': result.get('evidence', 'Must mention management company name and status')
                })

        elif 'risk' in rule_text_lower:
            result = llm_check_structural_element(section_text, 'risk_profile')

            if not result.get('found') or not result.get('appears_comprehensive') or result.get('confidence', 0) < 70:
                violations.append({
                    'type': 'STRUCTURE',
                    'severity': severity,
                    'slide': section,
                    'location': json_key,
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Risk profile incomplete (confidence: {result.get('confidence', 0)}%)",
                    'evidence': 'Risk information must be comprehensive and match prospectus'
                })

    return violations
def llm_check_general_requirement(doc_text, requirement_type, metadata=None):
    """
    LLM-based general requirement checker with multilingual support

    Args:
        doc_text: Full document text
        requirement_type: Type of check (retail_disclaimers, prof_disclaimers, sources, etc.)
        metadata: Additional context

    Returns:
        dict with found, missing_elements, confidence, evidence
    """
    if not tokenfactory_client:
        return {'found': False, 'confidence': 50}

    prompts = {
        'retail_disclaimers': f"""Does this document contain the REQUIRED RETAIL DISCLAIMERS?

DOCUMENT TEXT (excerpt):
{doc_text[:2500]}

REQUIRED ELEMENTS:
1. Capital at risk warning
   - English: "capital at risk", "risk of capital loss", "capital not guaranteed", "you may lose your capital"
   - French: "risque de perte en capital", "capital non garanti", "perte en capital", "risque de perte"

2. Past performance warning
   - English: "past performance is not indicative of future results", "past performance not guaranteed", "historical performance"
   - French: "les performances passÃ©es ne prÃ©jugent pas des performances futures", "performances passÃ©es ne garantissent pas"

Check SEMANTIC meaning, not exact wording. Both elements must be present.

Respond JSON:
{{
  "capital_risk_found": true/false,
  "capital_risk_text": "exact phrase or null",
  "past_performance_found": true/false,
  "past_performance_text": "exact phrase or null",
  "both_present": true/false,
  "confidence": 0-100,
  "language_detected": "en/fr/both"
}}""",

        'prof_disclaimers': f"""Does this document contain PROFESSIONAL INVESTOR disclaimer?

DOCUMENT TEXT (excerpt):
{doc_text[:2000]}

Look for:
- English: "professional investors", "professional clients", "reserved for professional investors", "institutional investors only", "qualified investors", "eligible counterparties"
- French: "investisseurs professionnels", "clients professionnels", "rÃ©servÃ© aux investisseurs professionnels", "rÃ©servÃ© aux professionnels", "investisseurs qualifiÃ©s"

Must clearly restrict document to professional audience.

Respond JSON:
{{
  "found": true/false,
  "expression_found": "exact phrase",
  "is_restrictive": true/false,
  "confidence": 0-100
}}""",

        'sources_dates': f"""Does numerical data in this document have proper SOURCE and DATE citations?

DOCUMENT TEXT (excerpt):
{doc_text[:3000]}

RULES:
1. Market data, statistics, index values â†’ NEED source + date
2. Fund's OWN data (AUM, NAV, fees from prospectus) â†’ NO source needed
3. Prospectus data (SRI, allocation) â†’ NO source needed

Look for:
- English: "Source:", "as of", "as at", "data from", "Bloomberg", "Morningstar"
- French: "Source :", "au", "Ã  la date du", "donnÃ©es au", "arrÃªtÃ© au", "chiffres au"

Analyze if data points requiring sources actually have them.

Respond JSON:
{{
  "has_external_data": true/false,
  "has_source_citations": true/false,
  "missing_sources": ["list of data without sources"],
  "confidence": 0-100,
  "reasoning": "explanation"
}}""",

        'sri_indicator': f"""Does this document contain the SRI (Synthetic Risk Indicator)?

DOCUMENT TEXT (excerpt):
{doc_text[:2500]}

Look for:
- "SRI", "SRRI", "Synthetic Risk Indicator", "Synthetic Risk and Reward Indicator"
- French: "Indicateur de risque", "SRRI", "Indicateur synthÃ©tique de risque"
- Format: "X/7" or "X sur 7" where X is 1-7
- Risk level mentions with scale

Respond JSON:
{{
  "found": true/false,
  "sri_value": "X/7 or null",
  "expression_found": "exact text",
  "has_disclaimer_nearby": true/false,
  "confidence": 0-100
}}""",

        'glossary': f"""Does this RETAIL document have a GLOSSARY for technical terms?

DOCUMENT TEXT (excerpt):
{doc_text[:2000]} ... {doc_text[-1000:]}

TECHNICAL TERMS TO CHECK: alpha, beta, volatility, duration, derivatives, tracking error, sharpe ratio

Look for:
- English: "Glossary", "Definitions", "Terms Used", "Key Concepts"
- French: "Glossaire", "Lexique", "DÃ©finitions", "Terminologie"

Usually at end of document.

Respond JSON:
{{
  "has_technical_terms": true/false,
  "terms_found": ["list"],
  "has_glossary": true/false,
  "glossary_location": "description or null",
  "confidence": 0-100
}}""",

        'internal_limits': f"""Does this document mention INTERNAL LIMITS (which is PROHIBITED)?

DOCUMENT TEXT (excerpt):
{doc_text[:2500]}

PROHIBITED MENTIONS:
- English: "internal limit", "internal threshold", "risk committee decision", "internal VaR", "internal constraint", "proprietary limit"
- French: "limite interne", "seuil interne", "contrainte interne", "dÃ©cision du comitÃ© des risques", "VaR interne"

These are non-contractual and must NOT appear in client documents.

Respond JSON:
{{
  "found_prohibited": true/false,
  "prohibited_phrases": ["list of found phrases"],
  "locations": ["where found"],
  "confidence": 0-100
}}""",

        'etf_liquidity': f"""Does this document make PROHIBITED liquidity claims about ETFs?

DOCUMENT TEXT (excerpt):
{doc_text[:2000]}

PROHIBITED CLAIMS:
- English: "liquid ETF", "high liquidity", "guaranteed liquidity", "easily sellable", "highly tradable", "liquid investment"
- French: "ETF liquide", "forte liquiditÃ©", "liquiditÃ© garantie", "facilement nÃ©gociable", "trÃ¨s liquide", "liquiditÃ© Ã©levÃ©e"

Cannot state ETFs are liquid or have guaranteed liquidity.

Respond JSON:
{{
  "document_mentions_etf": true/false,
  "has_liquidity_claims": true/false,
  "prohibited_phrases": ["list"],
  "confidence": 0-100
}}""",

        'morningstar_date': f"""If document shows Morningstar rating, does it include calculation DATE?

DOCUMENT TEXT (excerpt):
{doc_text[:2000]}

Check for:
1. Morningstar mention or stars (â˜…)
2. Date nearby with formats:
   - English: "as of DD/MM/YYYY", "as at Month YYYY"
   - French: "au DD/MM/YYYY", "Ã  la date du", "arrÃªtÃ© au"

Respond JSON:
{{
  "has_morningstar_rating": true/false,
  "has_date": true/false,
  "date_found": "exact text or null",
  "confidence": 0-100
}}""",

        'morningstar_category': f"""If document shows Morningstar rating, does it include reference CATEGORY?

DOCUMENT TEXT (excerpt):
{doc_text[:2000]}

Check for:
- English: "category", "EAA Fund", "peer group", "in the [Category Name] category"
- French: "catÃ©gorie", "univers", "dans la catÃ©gorie"

Must specify which Morningstar category the rating is from.

Respond JSON:
{{
  "has_morningstar_rating": true/false,
  "has_category": true/false,
  "category_found": "exact text or null",
  "confidence": 0-100
}}""",

        'team_disclaimer': f"""If document presents management team, does it have "subject to change" disclaimer?

DOCUMENT TEXT (excerpt):
{doc_text[:2500]}

Check for:
1. Team presentation (managers, team members)
2. Disclaimer nearby:
   - English: "subject to change", "may change", "composition may evolve", "not guaranteed"
   - French: "susceptible d'Ã©voluer", "peut changer", "composition non garantie", "peut Ãªtre modifiÃ©e"

Respond JSON:
{{
  "presents_team": true/false,
  "has_disclaimer": true/false,
  "disclaimer_text": "exact phrase or null",
  "confidence": 0-100
}}"""
    }

    prompt = prompts.get(requirement_type)
    if not prompt:
        return {'found': False, 'confidence': 0}

    result = call_tokenfactory_with_debug(
        prompt=prompt,
        system_message="You are a multilingual financial compliance expert. Understand semantic meaning, context, and different expressions in English and French. Respond ONLY with valid JSON.",
        function_name=f"check_general_requirement ({requirement_type})",
        show_prompt=True
    )
    
    if result:
        return result
    else:
        print(f"    âš ï¸  LLM error for {requirement_type}: Failed to parse response")
        return {'found': False, 'confidence': 50, 'error': 'JSON parsing failed'}


def check_general_rules_enhanced(doc, client_type, country_code=None):
    """
    Enhanced general rules checker with LLM multilingual validation
    """
    violations = []

    if not general_rules:
        return violations

    doc_text = extract_all_text_from_doc(doc)

    for rule in general_rules:
        rule_id = rule.get('rule_id', '')
        severity = rule.get('severity', 'major').upper()
        applies_to = rule.get('applies_to', {})

        # Check applicability
        if client_type.lower() not in [ct.lower() for ct in applies_to.get('client_type', [])]:
            continue

        if 'country' in applies_to and country_code and country_code not in applies_to['country']:
            continue

        # GEN_001: Retail disclaimers
        if rule_id == 'GEN_001' and client_type.lower() == 'retail':
            result = llm_check_general_requirement(doc_text, 'retail_disclaimers')

            if not result.get('both_present') or result.get('confidence', 0) < 70:
                missing = []
                if not result.get('capital_risk_found'):
                    missing.append('Capital at risk warning')
                if not result.get('past_performance_found'):
                    missing.append('Past performance warning')

                violations.append({
                    'type': 'GENERAL',
                    'severity': severity,
                    'slide': 'Document-wide',
                    'location': 'Required disclaimers',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Required retail disclaimers missing or incomplete: {', '.join(missing)}",
                    'evidence': f"Confidence: {result.get('confidence', 0)}%. Language: {result.get('language_detected', 'unknown')}"
                })

        # GEN_002: Professional disclaimers
        elif rule_id == 'GEN_002' and client_type.lower() == 'professional':
            result = llm_check_general_requirement(doc_text, 'prof_disclaimers')

            if not result.get('found') or not result.get('is_restrictive') or result.get('confidence', 0) < 70:
                violations.append({
                    'type': 'GENERAL',
                    'severity': severity,
                    'slide': 'Document-wide',
                    'location': 'Missing disclaimer',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': 'Professional client disclaimer missing or not sufficiently restrictive',
                    'evidence': f"Must clearly state: 'Reserved for professional investors'. Confidence: {result.get('confidence', 0)}%"
                })

        # GEN_003: Sources and dates
        elif rule_id == 'GEN_003':
            result = llm_check_general_requirement(doc_text, 'sources_dates')

            if result.get('has_external_data') and not result.get('has_source_citations'):
                violations.append({
                    'type': 'GENERAL',
                    'severity': severity,
                    'slide': 'Data sections',
                    'location': 'Multiple locations',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"External data without proper source/date citations",
                    'evidence': f"Missing sources for: {', '.join(result.get('missing_sources', ['various data points'])[:3])}. Confidence: {result.get('confidence', 0)}%"
                })

        # GEN_004: SRI requirement
        elif rule_id == 'GEN_004':
            result = llm_check_general_requirement(doc_text, 'sri_indicator')

            if not result.get('found') or result.get('confidence', 0) < 70:
                violations.append({
                    'type': 'GENERAL',
                    'severity': severity,
                    'slide': 'Fund presentation',
                    'location': 'Missing',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': 'SRI (Synthetic Risk Indicator) missing from fund presentation',
                    'evidence': f"Must include SRI X/7 with disclaimer on same slide. Confidence: {result.get('confidence', 0)}%"
                })

        # GEN_005: Glossary for retail
        elif rule_id == 'GEN_005' and client_type.lower() == 'retail':
            result = llm_check_general_requirement(doc_text, 'glossary')

            if result.get('has_technical_terms') and not result.get('has_glossary'):
                violations.append({
                    'type': 'GENERAL',
                    'severity': severity,
                    'slide': 'End of document',
                    'location': 'Missing glossary',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Technical terms used without glossary: {', '.join(result.get('terms_found', [])[:5])}",
                    'evidence': f"Retail documents must include glossary. Confidence: {result.get('confidence', 0)}%"
                })

        # GEN_012: Internal limits prohibition
        elif rule_id == 'GEN_012':
            result = llm_check_general_requirement(doc_text, 'internal_limits')

            if result.get('found_prohibited'):
                violations.append({
                    'type': 'GENERAL',
                    'severity': severity,
                    'slide': 'Document-wide',
                    'location': ', '.join(result.get('locations', ['Multiple sections'])),
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Internal limits mentioned (PROHIBITED): {', '.join(result.get('prohibited_phrases', []))}",
                    'evidence': f"Internal limits are non-contractual and must not appear. Confidence: {result.get('confidence', 0)}%"
                })

        # GEN_018: ETF liquidity claims
        elif rule_id == 'GEN_018':
            result = llm_check_general_requirement(doc_text, 'etf_liquidity')

            if result.get('document_mentions_etf') and result.get('has_liquidity_claims'):
                violations.append({
                    'type': 'GENERAL',
                    'severity': severity,
                    'slide': 'ETF sections',
                    'location': 'Multiple locations',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Prohibited ETF liquidity claims: {', '.join(result.get('prohibited_phrases', []))}",
                    'evidence': f"Cannot state ETF is liquid or has guaranteed liquidity. Confidence: {result.get('confidence', 0)}%"
                })

        # GEN_021: Morningstar date
        elif rule_id == 'GEN_021':
            result = llm_check_general_requirement(doc_text, 'morningstar_date')

            if result.get('has_morningstar_rating') and not result.get('has_date'):
                violations.append({
                    'type': 'GENERAL',
                    'severity': severity,
                    'slide': 'Morningstar section',
                    'location': 'Missing date',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': 'Morningstar rating without calculation date',
                    'evidence': f"Must include: 'as of DD/MM/YYYY'. Confidence: {result.get('confidence', 0)}%"
                })

        # GEN_022: Morningstar category
        elif rule_id == 'GEN_022':
            result = llm_check_general_requirement(doc_text, 'morningstar_category')

            if result.get('has_morningstar_rating') and not result.get('has_category'):
                violations.append({
                    'type': 'GENERAL',
                    'severity': severity,
                    'slide': 'Morningstar section',
                    'location': 'Missing category',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': 'Morningstar rating without reference category',
                    'evidence': f"Must specify category (e.g., 'in the Europe Equities category'). Confidence: {result.get('confidence', 0)}%"
                })

        # GEN_027: Team disclaimer
        elif rule_id == 'GEN_027':
            result = llm_check_general_requirement(doc_text, 'team_disclaimer')

            if result.get('presents_team') and not result.get('has_disclaimer'):
                violations.append({
                    'type': 'GENERAL',
                    'severity': severity,
                    'slide': 'Management team page',
                    'location': 'Missing disclaimer',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': 'Management team without "subject to change" disclaimer',
                    'evidence': f"Must add disclaimer. Confidence: {result.get('confidence', 0)}%"
                })

    return violations
def llm_analyze_security_mention(security_name, contexts, rule_type='recommendation'):
    """
    Analyze security mentions for investment recommendations or repetition context

    Args:
        security_name: Name of security
        contexts: List of text contexts where security appears
        rule_type: 'recommendation' or 'repetition'

    Returns:
        dict with analysis results
    """
    if not tokenfactory_client:
        return {'violation_found': False, 'confidence': 50}

    if rule_type == 'recommendation':
        contexts_text = '\n\n'.join([f"Context {i+1}:\n{ctx}" for i, ctx in enumerate(contexts[:5])])

        prompt = f"""Analyze if these mentions of "{security_name}" contain INVESTMENT RECOMMENDATIONS.

{contexts_text}

PROHIBITED LANGUAGE (Investment advice under MAR regulation):
- English: "we recommend", "buy", "sell", "strong buy", "accumulate", "target price", "undervalued", "overvalued", "attractive valuation", "should invest", "opportunity to buy"
- French: "nous recommandons", "acheter", "vendre", "achat fort", "accumuler", "prix cible", "cours cible", "sous-Ã©valuÃ©", "sur-Ã©valuÃ©", "valorisation attractive", "opportunitÃ© d'achat"

ALLOWED LANGUAGE (Factual descriptions):
- "The fund holds X shares"
- "Top holding: X"
- "X represents Y% of portfolio"
- "X operates in sector Y"

Question: Do ANY of these contexts contain investment recommendation language?

Respond JSON:
{{
  "contains_recommendation": true/false,
  "recommendation_phrases": ["list of prohibited phrases found"],
  "contexts_with_violations": [1, 2, ...],
  "confidence": 0-100,
  "reasoning": "explanation"
}}"""

    else:  # repetition analysis
        if len(contexts) < 2:
            return {'distinct_contexts': True, 'confidence': 100}

        contexts_text = '\n\n'.join([f"Mention {i+1}:\n{ctx}" for i, ctx in enumerate(contexts[:5])])

        prompt = f"""Analyze if these {len(contexts)} mentions of "{security_name}" are in DISTINCT contexts.

{contexts_text}

Questions:
1. Are these mentions in DIFFERENT contexts (e.g., holdings list vs. portfolio breakdown vs. risk analysis)?
2. Or are they REDUNDANT mentions of the same information?

DISTINCT contexts (OK):
- Mention 1: In top holdings list
- Mention 2: In sector breakdown chart
- Mention 3: In risk analysis
This is acceptable - different analytical purposes

REDUNDANT mentions (VIOLATION):
- Mention 1: "Top holding: Apple"
- Mention 2: "We like Apple's fundamentals"
- Mention 3: "Apple is well-positioned"
This violates the rule - same security highlighted multiple times

Respond JSON:
{{
  "are_distinct_contexts": true/false,
  "contexts_identified": ["list of context types for each mention"],
  "appears_redundant": true/false,
  "confidence": 0-100,
  "reasoning": "explanation"
}}"""

    result = call_tokenfactory_with_debug(
        prompt=prompt,
        system_message="You are a financial compliance expert specializing in MAR (Market Abuse Regulation). Analyze context carefully to detect investment recommendations. Respond ONLY with valid JSON.",
        function_name=f"analyze_security_mention ({security_name}, {rule_type})",
        show_prompt=True,
        max_tokens=800
    )
    
    if result:
        return result
    else:
        print(f"    âš ï¸  LLM error analyzing {security_name}: Failed to parse response")
        return {'violation_found': False, 'confidence': 50, 'error': 'JSON parsing failed'}


def extract_security_contexts(doc, security_name, context_chars=400):
    """Extract text contexts around each security mention"""
    doc_text = extract_all_text_from_doc(doc)
    doc_lower = doc_text.lower()
    security_lower = security_name.lower()

    contexts = []
    start = 0
    while True:
        pos = doc_lower.find(security_lower, start)
        if pos == -1:
            break

        context_start = max(0, pos - context_chars)
        context_end = min(len(doc_text), pos + len(security_name) + context_chars)
        context = doc_text[context_start:context_end]
        contexts.append(context)

        start = pos + 1

    return contexts


def llm_check_prohibited_phrases(doc_text, rule_id, prohibited_phrases_en, prohibited_phrases_fr=None):
    """
    Check for prohibited phrases with LLM understanding semantic equivalents

    Args:
        doc_text: Document text
        rule_id: Rule identifier
        prohibited_phrases_en: List of English prohibited phrases
        prohibited_phrases_fr: List of French prohibited phrases

    Returns:
        dict with violations found
    """
    if not tokenfactory_client:
        # Fallback to simple matching
        found = []
        for phrase in prohibited_phrases_en:
            if phrase.lower() in doc_text.lower():
                found.append(phrase)
        return {'violations_found': found, 'confidence': 60}

    prohibited_phrases_fr = prohibited_phrases_fr or []

    prompt = f"""Analyze if this document contains INVESTMENT RECOMMENDATION language (PROHIBITED under MAR).

DOCUMENT TEXT (excerpt):
{doc_text[:3000]}

PROHIBITED PHRASES - English:
{', '.join(prohibited_phrases_en[:15])}

PROHIBITED PHRASES - French:
{', '.join(prohibited_phrases_fr[:15]) if prohibited_phrases_fr else 'N/A'}

Look for:
1. Exact matches of these phrases
2. SEMANTIC EQUIVALENTS in English or French
3. Variations or synonyms with same meaning

Examples:
- "We recommend buying" = VIOLATION
- "Strong investment opportunity" = VIOLATION
- "Nous recommandons l'achat" = VIOLATION
- "Valorisation attractive pour investir" = VIOLATION

DO NOT flag factual descriptions:
- "The fund invests in..." = OK
- "Holdings include..." = OK

Respond JSON:
{{
  "violations_found": ["list of actual phrases found in document"],
  "locations": ["brief description of where found"],
  "confidence": 0-100
}}"""

    result = call_tokenfactory_with_debug(
        prompt=prompt,
        system_message="You are a MAR compliance expert. Detect investment recommendations in any language. Respond ONLY with valid JSON.",
        function_name=f"check_prohibited_phrases ({rule_id})",
        show_prompt=True
    )
    
    if result:
        return result
    else:
        print(f"    âš ï¸  LLM error checking phrases for {rule_id}: Failed to parse response")
        return {'violations_found': [], 'confidence': 50}

def extract_security_mentions(doc):
    """Extract and count mentions of specific securities/companies"""
    # This is a simplified version - in production, you'd use NER or a companies database
    doc_text = extract_all_text_from_doc(doc)

    # Common patterns for company names (simplified)
    # In production, use a proper entity recognition system
    company_patterns = [
        r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+(?:Inc|Corp|Ltd|SA|SAS|AG|GmbH|PLC)\b',
        r'\b([A-Z]{2,})\b'  # Acronyms like AAPL, MSFT
    ]

    companies = []
    for pattern in company_patterns:
        matches = re.findall(pattern, doc_text)
        companies.extend(matches)

    # Count occurrences
    company_counts = Counter(companies)

    # Filter out common words that might be false positives
    common_words = {'THE', 'AND', 'FOR', 'ARE', 'THIS', 'WITH', 'FROM', 'THAT', 'HAVE', 'BEEN'}
    filtered_counts = {k: v for k, v in company_counts.items() if k not in common_words and len(k) > 2}

    return filtered_counts


def check_repeated_securities_ai(doc):
    """
    AI-enhanced check for repeated external company mentions with whitelist filtering
    
    Eliminates 16 false positives by whitelisting:
    - Fund name components (e.g., "ODDO", "BHF")
    - Strategy terminology (e.g., "momentum", "quantitative")
    - Regulatory terms (e.g., "SRI", "SRRI", "SFDR")
    - Benchmark names (e.g., "S&P", "MSCI")
    - Generic financial terms (e.g., "fund", "investment")
    
    Only flags external company names mentioned 3+ times after whitelist filtering.
    Uses SemanticValidator to verify if terms are actually company names.
    
    Args:
        doc: Document dictionary
    
    Returns:
        List of violations (empty if no violations found)
    
    Test cases:
        - "ODDO BHF" (31 mentions) → should NOT flag (fund name)
        - "momentum" (2 mentions) → should NOT flag (strategy term)
        - "SRI" (2 mentions) → should NOT flag (regulatory term)
        - "Apple" (5 mentions) → should flag (external company, 3+ mentions)
    """
    violations = []
    
    try:
        import logging
        logger = logging.getLogger(__name__)
        
        from ai_engine import create_ai_engine_from_env
        from whitelist_manager import WhitelistManager
        from semantic_validator import SemanticValidator
        
        # Initialize components
        whitelist_mgr = WhitelistManager()
        
        # Build whitelist from document metadata
        whitelist = whitelist_mgr.build_whitelist(doc)
        
        logger.info(f"Built whitelist with {len(whitelist)} terms")
        logger.info(f"Fund name terms: {whitelist_mgr.fund_name_terms}")
        
        # Extract all text from document
        doc_text = extract_all_text_from_doc(doc)
        
        if not doc_text or len(doc_text) < 50:
            return violations
        
        # Extract capitalized words as potential company names
        # Pattern: Words that start with capital letter (2+ chars)
        capitalized_words = re.findall(r'\b[A-Z][A-Za-z]+\b', doc_text)
        
        # Also extract acronyms (2-5 uppercase letters)
        acronyms = re.findall(r'\b[A-Z]{2,5}\b', doc_text)
        
        # Combine and count
        all_terms = capitalized_words + acronyms
        term_counts = Counter(term.lower() for term in all_terms)
        
        logger.info(f"Found {len(term_counts)} unique capitalized terms/acronyms")
        
        # Filter out whitelisted terms BEFORE counting threshold
        non_whitelisted_terms = {}
        for term, count in term_counts.items():
            if not whitelist_mgr.is_whitelisted(term):
                non_whitelisted_terms[term] = count
            else:
                reason = whitelist_mgr.get_whitelist_reason(term)
                logger.info(f"Skipping whitelisted term: {term} ({count} mentions) - {reason}")
        
        logger.info(f"After whitelist filtering: {len(non_whitelisted_terms)} terms remain")
        
        # Only check terms mentioned 3+ times
        repeated_terms = {term: count for term, count in non_whitelisted_terms.items() if count >= 3}
        
        if not repeated_terms:
            logger.info("No repeated non-whitelisted terms found (threshold: 3+ mentions)")
            return violations
        
        logger.info(f"Found {len(repeated_terms)} terms with 3+ mentions: {list(repeated_terms.keys())}")
        
        # Initialize AI engine for semantic validation
        ai_engine = create_ai_engine_from_env()
        
        if not ai_engine:
            logger.warning("AI engine not available, using rule-based validation only")
            # Fallback: flag all repeated non-whitelisted terms
            for term, count in repeated_terms.items():
                violations.append({
                    'type': 'SECURITIES/VALUES',
                    'severity': 'MAJOR',
                    'slide': 'Multiple slides',
                    'location': 'Document-wide',
                    'rule': 'VAL_005',
                    'message': f'External company "{term}" mentioned {count} times',
                    'evidence': f'Term appears {count} times in document (rule-based detection, no AI verification)',
                    'confidence': 60
                })
            return violations
        
        # Use SemanticValidator to verify each term
        semantic_validator = SemanticValidator(ai_engine)
        
        # Limit to top 10 most mentioned terms for performance
        top_repeated = sorted(repeated_terms.items(), key=lambda x: x[1], reverse=True)[:10]
        
        for term, count in top_repeated:
            logger.info(f"Validating term: {term} ({count} mentions)")
            
            # Use SemanticValidator to check if it's actually an external company
            result = semantic_validator.validate_securities_mention(
                text=doc_text,
                term=term,
                whitelist=whitelist,
                mention_count=count
            )
            
            if result.is_violation:
                violations.append({
                    'type': 'SECURITIES/VALUES',
                    'severity': 'MAJOR',
                    'slide': 'Multiple slides',
                    'location': 'Document-wide',
                    'rule': 'VAL_005',
                    'message': f'External company "{term}" mentioned {count} times',
                    'evidence': f'{result.reasoning}. {result.evidence[0] if result.evidence else ""}',
                    'confidence': result.confidence,
                    'method': result.method
                })
                logger.info(f"  → VIOLATION: {result.reasoning} (confidence: {result.confidence}%)")
            else:
                logger.info(f"  → OK: {result.reasoning} (confidence: {result.confidence}%)")
        
        return violations
    
    except Exception as e:
        import logging
        import traceback
        logger = logging.getLogger(__name__)
        logger.error(f"Error in check_repeated_securities_ai: {e}")
        traceback.print_exc()
        return violations
def check_values_rules_enhanced(doc):
    """
    Enhanced securities/values rules checker with LLM context analysis
    """
    violations = []

    if not values_rules:
        return violations

    doc_text = extract_all_text_from_doc(doc)

    # Check for investment advice using NEW AI context-aware function
    # This eliminates 25 false positives by understanding context
    # Only check once for all prohibition rules (not per rule)
    prohibition_rules = [r for r in values_rules if r.get('category') != 'allowed']
    
    print(f"\n[DEBUG] Found {len(prohibition_rules)} prohibition rules")
    
    if prohibition_rules:
        # Use the first prohibition rule as representative (they all check for investment advice)
        representative_rule = prohibition_rules[0]
        print(f"[DEBUG] Calling check_prohibited_phrases_ai with rule {representative_rule.get('rule_id')}")
        
        try:
            ai_violations = check_prohibited_phrases_ai(doc, representative_rule)
            print(f"[DEBUG] AI check returned {len(ai_violations)} violations")
            
            if ai_violations:
                violations.extend(ai_violations)
        except Exception as e:
            print(f"[ERROR] AI check failed: {e}")
            import traceback
            traceback.print_exc()

    # Check for repeated security mentions with NEW whitelist-aware AI function
    # This eliminates 16 false positives by whitelisting fund names, strategy terms, etc.
    print("\n🔍 Analyzing repeated security mentions with whitelist-aware AI...")
    
    try:
        repeated_violations = check_repeated_securities_ai(doc)
        print(f"[DEBUG] Whitelist-aware check returned {len(repeated_violations)} violations")
        
        if repeated_violations:
            violations.extend(repeated_violations)
    except Exception as e:
        print(f"[ERROR] Whitelist-aware check failed: {e}")
        import traceback
        traceback.print_exc()

    return violations


def check_prohibited_phrases_ai(doc, rule):
    """
    AI context-aware check for prohibited investment advice phrases
    
    Eliminates 25 false positives by distinguishing:
    - Fund strategy descriptions (ALLOWED): "Le fonds investit dans...", "Tirer parti du momentum"
    - Investment advice to clients (PROHIBITED): "Vous devriez investir", "Nous recommandons"
    
    Uses ContextAnalyzer and IntentClassifier to understand:
    - WHO is performing the action (fund vs client)
    - WHAT is the intent (describe vs advise)
    
    Args:
        doc: Document dictionary
        rule: Rule dictionary with prohibited phrases
    
    Returns:
        List of violations (empty if no violations found)
    """
    violations = []
    
    try:
        # Import AI components
        import logging
        logger = logging.getLogger(__name__)
        
        from ai_engine import create_ai_engine_from_env
        from context_analyzer import ContextAnalyzer
        from intent_classifier import IntentClassifier
        
        # Initialize AI engine
        ai_engine = create_ai_engine_from_env()
        if not ai_engine:
            logger.warning("AI engine not available, skipping AI context-aware check")
            return violations
        
        # Initialize components
        context_analyzer = ContextAnalyzer(ai_engine)
        intent_classifier = IntentClassifier(ai_engine)
        
        # Extract all text from document
        doc_text = extract_all_text_from_doc(doc)
        
        if not doc_text or len(doc_text) < 50:
            return violations
        
        # Get rule details
        rule_id = rule.get('rule_id', 'VAL_001')
        severity = rule.get('severity', 'CRITICAL').upper()
        prohibited_phrases = rule.get('prohibited_phrases', [])
        
        # Quick check: if no prohibited phrases in text, skip AI analysis
        text_lower = doc_text.lower()
        has_potential_violation = False
        
        for phrase in prohibited_phrases[:10]:  # Check first 10 phrases
            if phrase.lower() in text_lower:
                has_potential_violation = True
                break
        
        if not has_potential_violation:
            # No prohibited phrases found, no violation
            return violations
        
        # Analyze context with AI
        print(f"\n[AI] Analyzing context for {rule_id} with AI...")
        
        # Use ContextAnalyzer to understand the text
        context_analysis = context_analyzer.analyze_context(doc_text, "investment_advice")
        
        # Use IntentClassifier to determine intent
        intent_classification = intent_classifier.classify_intent(doc_text)
        
        # Decision logic: Only flag if BOTH conditions are met:
        # 1. Intent is ADVICE (not DESCRIPTION)
        # 2. Subject is CLIENT (not FUND)
        
        is_client_advice = (
            intent_classification.intent_type == "ADVICE" and
            context_analysis.subject == "client"
        )
        
        # Alternative check using helper methods
        is_advice_alt = intent_classifier.is_client_advice(doc_text)
        is_fund_desc = context_analyzer.is_fund_strategy_description(doc_text)
        
        # Combine results with confidence scoring
        combined_confidence = min(
            context_analysis.confidence,
            intent_classification.confidence
        )
        
        # Only flag if high confidence client advice
        if (is_client_advice or is_advice_alt) and not is_fund_desc and combined_confidence >= 70:
            # Extract evidence of prohibited phrases
            evidence_phrases = []
            for phrase in prohibited_phrases[:5]:
                if phrase.lower() in text_lower:
                    # Find context around phrase
                    idx = text_lower.find(phrase.lower())
                    if idx != -1:
                        start = max(0, idx - 50)
                        end = min(len(doc_text), idx + len(phrase) + 50)
                        context = doc_text[start:end].strip()
                        evidence_phrases.append(f'"{phrase}" in context: ...{context}...')
            
            violations.append({
                'type': 'SECURITIES/VALUES',
                'severity': severity,
                'slide': 'Multiple locations',
                'location': 'Document-wide',
                'rule': f"{rule_id}: {rule.get('rule_text', 'No investment recommendations')}",
                'message': f"Investment advice to clients detected (MAR violation)",
                'evidence': (
                    f"AI Analysis:\n"
                    f"  - Intent: {intent_classification.intent_type} (confidence: {intent_classification.confidence}%)\n"
                    f"  - Subject: {context_analysis.subject} (confidence: {context_analysis.confidence}%)\n"
                    f"  - Context Reasoning: {context_analysis.reasoning}\n"
                    f"  - Intent Reasoning: {intent_classification.reasoning}\n"
                    f"  - Evidence: {', '.join(evidence_phrases[:3]) if evidence_phrases else 'See prohibited phrases'}\n"
                    f"  - Method: AI_CONTEXT_AWARE"
                )
            })
            
            print(f"  [X] Violation detected: Client advice (confidence: {combined_confidence}%)")
        else:
            # Not a violation - likely fund description
            print(f"  [OK] No violation: Fund description detected")
            print(f"     - Intent: {intent_classification.intent_type}")
            print(f"     - Subject: {context_analysis.subject}")
            print(f"     - Is fund description: {is_fund_desc}")
            print(f"     - Confidence: {combined_confidence}%")
    
    except Exception as e:
        logger.error(f"AI context-aware check failed: {e}")
        import traceback
        traceback.print_exc()
    
    return violations


def llm_detect_performance_content(slide_text, slide_number, metadata=None):
    """
    Use LLM to detect performance data including charts, tables, implicit mentions

    Args:
        slide_text: Full text from slide
        slide_number: Slide position
        metadata: Additional context

    Returns:
        dict with shows_performance, performance_type, confidence
    """
    if not tokenfactory_client:
        # Fallback to keyword matching
        perf_keywords = ['performance', 'return', 'yield', '%', 'ytd']
        has_perf = any(kw in slide_text.lower() for kw in perf_keywords)
        return {'shows_performance': has_perf, 'confidence': 60}

    # Check for numerical patterns that suggest performance
    has_percentages = bool(re.findall(r'\d+\.?\d*\s*%', slide_text))
    has_time_labels = bool(re.search(r'\b(ytd|1y|3y|5y|10y|year|year|an|ans)\b', slide_text, re.I))
    has_large_numbers = bool(re.findall(r'\+\d+%|\-\d+%', slide_text))

    prompt = f"""Analyze if this slide shows PERFORMANCE data (returns, results, fund evolution).

SLIDE {slide_number} CONTENT:
{slide_text[:1500]}

INDICATORS DETECTED:
- Has percentages: {has_percentages}
- Has time period labels: {has_time_labels}
- Has +/- percentage values: {has_large_numbers}

PERFORMANCE INDICATORS TO CHECK:
- English: "performance", "return", "yield", "growth", "gain", "evolution", "results", "cumulative", "annualized", "since inception", "+X%", "YTD"
- French: "performance", "rendement", "Ã©volution", "croissance", "progression", "rÃ©sultats", "depuis crÃ©ation", "+X%"

Also check for:
1. CHARTS/GRAPHS showing growth/returns over time
2. TABLES with performance figures
3. IMPLICIT performance: "Fund up 15%" without using word "performance"
4. Comparison lines, historical data

DO NOT flag:
- Asset allocation percentages (e.g., "60% equities")
- Fees percentages (e.g., "0.75% management fee")
- Risk indicator (e.g., "SRI: 4/7")

Question: Does this slide show fund PERFORMANCE/RETURNS data?

Respond JSON:
{{
  "shows_performance": true/false,
  "performance_type": "explicit" or "chart" or "table" or "implicit" or null,
  "evidence": "what indicates performance",
  "confidence": 0-100,
  "is_first_slide": {slide_number <= 2}
}}"""

    try:
        response = tokenfactory_client.chat.completions.create(
            model="hosted_vllm/Llama-3.1-70B-Instruct",
            messages=[
                {"role": "system", "content": "You are a financial document analyst. Detect performance data in any form: text, charts, tables, or implicit mentions. Respond ONLY with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=500
        )

        result_text = response.choices[0].message.content.strip()
        result_text = result_text.replace('```json', '').replace('```', '').strip()
        result = json.loads(result_text)

        return result
    except Exception as e:
        print(f"    âš ï¸  LLM error detecting performance on slide {slide_number}: {e}")
        return {'shows_performance': False, 'confidence': 50}


def llm_check_performance_context(doc_text, check_type, metadata=None):
    """
    Check performance-related requirements with LLM

    Args:
        doc_text: Document text
        check_type: 'ytd_context', 'benchmark', 'disclaimer', 'long_term'
        metadata: Additional context

    Returns:
        dict with check results
    """
    if not tokenfactory_client:
        return {'requirement_met': False, 'confidence': 50}

    prompts = {
        'ytd_context': f"""Does this document show YTD performance WITH proper long-term context?

DOCUMENT TEXT (excerpt):
{doc_text[:2500]}

RULES:
- If YTD performance shown â†’ Must ALSO show longer-term performance (1Y, 3Y, 5Y, 10Y, or since inception)
- English: "YTD", "year-to-date", "1 year", "3 years", "5 years", "10 years", "since inception", "since launch"
- French: "sur l'annÃ©e", "1 an", "3 ans", "5 ans", "10 ans", "depuis crÃ©ation", "depuis lancement"

Question: If YTD is shown, is it accompanied by long-term performance?

Respond JSON:
{{
  "shows_ytd": true/false,
  "ytd_text": "exact phrase or null",
  "has_long_term_context": true/false,
  "long_term_periods": ["list like '3Y', '5Y'"],
  "confidence": 0-100
}}""",

        'benchmark': f"""Does performance data include proper BENCHMARK comparison?

DOCUMENT TEXT (excerpt):
{doc_text[:2500]}

REQUIREMENT: Performance must be compared to official benchmark from prospectus.

Look for:
- English: "benchmark", "vs", "versus", "compared to", "reference index", "comparison"
- French: "indice de rÃ©fÃ©rence", "par rapport Ã ", "vs", "comparÃ© Ã ", "rÃ©fÃ©rence"

Question: Is performance compared to a benchmark?

Respond JSON:
{{
  "has_benchmark_comparison": true/false,
  "benchmark_mentioned": "name or null",
  "comparison_phrases": ["list"],
  "confidence": 0-100
}}""",

        'disclaimer': f"""Does performance section have the MANDATORY disclaimer?

DOCUMENT TEXT (excerpt):
{doc_text[:2500]}

REQUIRED ELEMENTS (must have ALL):
1. Past performance mention
2. Not indicative / not guaranteed
3. Future performance / future results

English variations:
- "Past performance is not indicative of future results"
- "Past performance does not guarantee future performance"
- "Historical performance is no guarantee of future returns"

French variations:
- "Les performances passÃ©es ne prÃ©jugent pas des performances futures"
- "Les performances passÃ©es ne garantissent pas les performances futures"

Question: Is the complete mandatory disclaimer present?

Respond JSON:
{{
  "has_complete_disclaimer": true/false,
  "disclaimer_text": "exact text found or null",
  "missing_elements": ["list of missing parts"],
  "confidence": 0-100
}}""",

        'long_term': f"""Does document show LONG-TERM performance (not just short-term)?

DOCUMENT TEXT (excerpt):
{doc_text[:2500]}

REQUIREMENT: Must show performance over meaningful periods (1Y minimum, preferably 3Y, 5Y, 10Y, or since inception)

Look for time periods:
- English: "1 year", "3 years", "5 years", "10 years", "since inception", "since launch"
- French: "1 an", "3 ans", "5 ans", "10 ans", "depuis crÃ©ation", "depuis lancement"

Question: Are long-term performance periods shown?

Respond JSON:
{{
  "has_long_term": true/false,
  "periods_found": ["list like '1Y', '3Y', '5Y'"],
  "longest_period": "period or null",
  "confidence": 0-100
}}"""
    }

    prompt = prompts.get(check_type)
    if not prompt:
        return {'requirement_met': False, 'confidence': 0}

    try:
        response = tokenfactory_client.chat.completions.create(
            model="hosted_vllm/Llama-3.1-70B-Instruct",
            messages=[
                {"role": "system", "content": "You are a financial compliance expert. Check performance presentation requirements. Understand context and multiple languages. Respond ONLY with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=600
        )

        result_text = response.choices[0].message.content.strip()
        result_text = result_text.replace('```json', '').replace('```', '').strip()
        result = json.loads(result_text)

        return result
    except Exception as e:
        print(f"    âš ï¸  LLM error checking {check_type}: {e}")
        return {'requirement_met': False, 'confidence': 50}


def check_performance_rules_enhanced(doc, client_type, fund_age_years=None):
    """
    Enhanced performance rules checker with LLM multi-modal detection
    """
    violations = []

    if not performance_rules:
        return violations

    doc_text = extract_all_text_from_doc(doc)

    # Detect performance content in each slide using LLM
    print("\nðŸ” Analyzing slides for performance content with LLM...")
    performance_slides = []

    # Check cover page
    if 'page_de_garde' in doc:
        cover_text = extract_section_text(doc, 'page_de_garde')
        result = llm_detect_performance_content(cover_text, 1)
        if result.get('shows_performance') and result.get('confidence', 0) > 70:
            performance_slides.append({
                'slide': 1,
                'section': 'Cover Page',
                'type': result.get('performance_type'),
                'evidence': result.get('evidence'),
                'confidence': result.get('confidence')
            })

    # Check other slides
    if 'pages_suivantes' in doc:
        for idx, slide in enumerate(doc['pages_suivantes'], start=3):
            slide_text = json.dumps(slide)
            slide_num = slide.get('slide_number', idx)
            result = llm_detect_performance_content(slide_text, slide_num)
            if result.get('shows_performance') and result.get('confidence', 0) > 70:
                performance_slides.append({
                    'slide': slide_num,
                    'section': f'Slide {slide_num}',
                    'type': result.get('performance_type'),
                    'evidence': result.get('evidence'),
                    'confidence': result.get('confidence')
                })

    has_performance = len(performance_slides) > 0

    if not has_performance:
        print("âœ… No performance content detected\n")
        return violations

    print(f"ðŸ“Š Performance detected on {len(performance_slides)} slide(s):")
    for perf in performance_slides[:5]:
        print(f"   â€¢ {perf['section']}: {perf['type']} (confidence: {perf['confidence']}%)")
    print()

    # Now check rules
    for rule in performance_rules:
        rule_id = rule.get('rule_id', '')
        severity = rule.get('severity', 'major').upper()
        applies_to = rule.get('applies_to', {})

        if client_type.lower() not in [ct.lower() for ct in applies_to.get('client_type', [])]:
            continue

        # PERF_001: Document can't start with performance
        if rule_id == 'PERF_001':
            first_slide_perf = [p for p in performance_slides if p['slide'] <= 2]
            if first_slide_perf:
                violations.append({
                    'type': 'PERFORMANCE',
                    'severity': severity,
                    'slide': first_slide_perf[0]['section'],
                    'location': 'Beginning of document',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f"Document starts with performance data ({first_slide_perf[0]['type']})",
                    'evidence': f"Detected: {first_slide_perf[0]['evidence']}. Confidence: {first_slide_perf[0]['confidence']}%. Performance must be preceded by fund presentation."
                })

        # PERF_010: YTD alone prohibited
        elif rule_id == 'PERF_010':
            result = llm_check_performance_context(doc_text, 'ytd_context')

            if result.get('shows_ytd') and not result.get('has_long_term_context'):
                violations.append({
                    'type': 'PERFORMANCE',
                    'severity': severity,
                    'slide': 'Performance section',
                    'location': 'YTD data',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': 'YTD performance shown without long-term context',
                    'evidence': f"Found YTD: '{result.get('ytd_text', 'N/A')}'. Must show 10Y/5Y/3Y/1Y or inception alongside YTD. Confidence: {result.get('confidence', 0)}%"
                })

        # PERF_011: Under 1 year = no performance
        elif rule_id == 'PERF_011':
            if fund_age_years is not None and fund_age_years < 1 and has_performance:
                violations.append({
                    'type': 'PERFORMANCE',
                    'severity': severity,
                    'slide': 'All performance slides',
                    'location': f"{len(performance_slides)} slide(s) with performance",
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': f'Fund is {fund_age_years:.1f} years old - cannot show any performance',
                    'evidence': 'Funds under 1 year cannot display performance under any circumstances'
                })

        # PERF_014: Must compare to benchmark
        elif rule_id == 'PERF_014':
            result = llm_check_performance_context(doc_text, 'benchmark')

            if not result.get('has_benchmark_comparison') or result.get('confidence', 0) < 70:
                violations.append({
                    'type': 'PERFORMANCE',
                    'severity': severity,
                    'slide': 'Performance section',
                    'location': 'All performance data',
                    'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                    'message': 'Performance shown without benchmark comparison',
                    'evidence': f"Must compare to official prospectus benchmark. Confidence: {result.get('confidence', 0)}%"
                })

        # PERF_054: Mandatory disclaimer
        elif rule_id == 'PERF_054':
            result = llm_check_performance_context(doc_text, 'disclaimer')

            if not result.get('has_complete_disclaimer') or result.get('confidence', 0) < 70:
                missing_parts = ', '.join(result.get('missing_elements', ['complete disclaimer']))
                violations.append({
                    'type': 'PERFORMANCE',
                    'severity': severity,
                    'slide': 'Near performance data',
                    'location': 'Missing or incomplete',
                    'rule': f"{rule_id}: Mandatory performance disclaimer",
                    'message': f'Required performance disclaimer missing or incomplete: {missing_parts}',
                    'evidence': f"Required: 'Past performance is not indicative of future results'. Confidence: {result.get('confidence', 0)}%"
                })

    return violations
def llm_analyze_esg_content(slide_text, slide_number):
    """
    Analyze if slide is primarily about ESG/sustainability

    Args:
        slide_text: Full slide text
        slide_number: Slide position

    Returns:
        dict with is_esg_slide, esg_percentage, content_type
    """
    if not tokenfactory_client:
        # Fallback to keyword counting
        esg_keywords = ['esg', 'environmental', 'social', 'governance', 'sustainability', 'sustainable']
        keyword_count = sum(1 for kw in esg_keywords if kw in slide_text.lower())
        is_esg = keyword_count >= 3
        return {'is_esg_slide': is_esg, 'esg_percentage': 80 if is_esg else 20, 'confidence': 60}

    prompt = f"""Analyze if this slide is PRIMARILY about ESG/Sustainability content.

SLIDE {slide_number} CONTENT:
{slide_text[:1500]}

ESG/SUSTAINABILITY INDICATORS:
- English: ESG, environmental, social, governance, sustainability, sustainable, responsible investment, SRI, green, climate, carbon, impact, exclusion, engagement, stewardship, SFDR, Article 8, Article 9, taxonomy, SDG
- French: ESG, environnemental, social, gouvernance, dÃ©veloppement durable, investissement responsable, ISR, vert, climat, carbone, impact, exclusion, engagement, SFDR, Article 8, Article 9, taxonomie, ODD

Question: What percentage of this slide's content is dedicated to ESG?

Categories:
- 0-20%: Brief mention, passing reference
- 21-50%: Significant mention but not primary focus
- 51-80%: Major focus, substantial ESG content
- 81-100%: Entirely dedicated to ESG

Respond JSON:
{{
  "is_esg_slide": true/false,
  "esg_percentage": 0-100,
  "content_type": "brief_mention" or "baseline_exclusions" or "esg_strategy" or "esg_detailed" or "non_esg",
  "esg_elements_found": ["list of ESG topics"],
  "confidence": 0-100
}}"""

    try:
        response = tokenfactory_client.chat.completions.create(
            model="hosted_vllm/Llama-3.1-70B-Instruct",
            messages=[
                {"role": "system", "content": "You are an ESG content analyst. Assess what percentage of content is ESG-related. Be precise in classification. Respond ONLY with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=500
        )

        result_text = response.choices[0].message.content.strip()
        result_text = result_text.replace('```json', '').replace('```', '').strip()
        result = json.loads(result_text)

        return result
    except Exception as e:
        print(f"    âš ï¸  LLM error analyzing ESG on slide {slide_number}: {e}")
        return {'is_esg_slide': False, 'esg_percentage': 0, 'confidence': 50}


def llm_check_esg_baseline(doc_text):
    """
    Check if ESG content goes beyond baseline exclusions

    Args:
        doc_text: Document text

    Returns:
        dict with beyond_baseline, esg_topics
    """
    if not tokenfactory_client:
        return {'beyond_baseline': True, 'confidence': 50}

    prompt = f"""Analyze if ESG content in this document goes BEYOND baseline exclusions.

DOCUMENT TEXT (excerpt):
{doc_text[:2500]}

BASELINE EXCLUSIONS (Allowed for all funds):
- English: "controversial weapons exclusion", "baseline exclusions", "firm-wide exclusions", "OBAM exclusion policy"
- French: "exclusion armes controversÃ©es", "exclusions socle", "exclusions de base", "politique d'exclusion OBAM"

BEYOND BASELINE (Not allowed for "Other" funds):
- Detailed ESG strategy, ESG integration methodology
- ESG scoring, ESG ratings, ESG analysis
- Climate strategy, carbon reduction targets
- Impact measurement, SDG alignment
- ESG engagement, stewardship activities
- SFDR classification details (Article 8/9)

Question: Does ESG content go beyond simple baseline exclusions?

Respond JSON:
{{
  "mentions_esg": true/false,
  "beyond_baseline": true/false,
  "baseline_only": true/false,
  "esg_topics_found": ["list of topics beyond baseline"],
  "confidence": 0-100,
  "reasoning": "explanation"
}}"""

    try:
        response = tokenfactory_client.chat.completions.create(
            model="hosted_vllm/Llama-3.1-70B-Instruct",
            messages=[
                {"role": "system", "content": "You are an ESG compliance expert. Distinguish between baseline exclusions and substantive ESG content. Respond ONLY with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=600
        )

        result_text = response.choices[0].message.content.strip()
        result_text = result_text.replace('```json', '').replace('```', '').strip()
        result = json.loads(result_text)

        return result
    except Exception as e:
        print(f"    âš ï¸  LLM error checking baseline: {e}")
        return {'beyond_baseline': True, 'confidence': 50}


def check_esg_rules_enhanced(doc, esg_classification, client_type):
    """
    Enhanced ESG rules checker with LLM section-based analysis
    """
    violations = []

    if not esg_rules:
        return violations

    doc_text = extract_all_text_from_doc(doc)
    classification_lower = esg_classification.lower()

    # Analyze each slide for ESG content
    print("\nðŸ” Analyzing ESG content distribution with LLM...")
    esg_slides = []
    total_chars = 0
    esg_chars = 0

    # Check all slides
    all_sections = []
    if 'page_de_garde' in doc:
        all_sections.append(('page_de_garde', 1))
    if 'slide_2' in doc:
        all_sections.append(('slide_2', 2))
    if 'pages_suivantes' in doc:
        for idx, slide in enumerate(doc['pages_suivantes'], start=3):
            slide_num = slide.get('slide_number', idx)
            all_sections.append((slide, slide_num))
    if 'page_de_fin' in doc:
        all_sections.append(('page_de_fin', 99))

    for section, slide_num in all_sections:
        if isinstance(section, str):
            slide_text = extract_section_text(doc, section)
        else:
            slide_text = json.dumps(section)

        slide_chars = len(slide_text)
        total_chars += slide_chars

        # Analyze with LLM
        result = llm_analyze_esg_content(slide_text, slide_num)

        if result.get('is_esg_slide') or result.get('esg_percentage', 0) > 20:
            esg_slides.append({
                'slide': slide_num,
                'esg_percentage': result.get('esg_percentage', 0),
                'content_type': result.get('content_type'),
                'elements': result.get('esg_elements_found', []),
                'chars': slide_chars,
                'confidence': result.get('confidence')
            })

            # Calculate ESG chars for this slide
            esg_chars += int(slide_chars * result.get('esg_percentage', 0) / 100)

    esg_volume_percentage = (esg_chars / total_chars * 100) if total_chars > 0 else 0

    print(f"ðŸ“Š ESG Content Analysis:")
    print(f"   â€¢ Total document: {total_chars:,} characters")
    print(f"   â€¢ ESG content: ~{esg_chars:,} characters ({esg_volume_percentage:.1f}%)")
    print(f"   â€¢ ESG slides: {len(esg_slides)}")
    if esg_slides:
        for slide_info in esg_slides[:5]:
            print(f"     - Slide {slide_info['slide']}: {slide_info['esg_percentage']}% ESG ({slide_info['content_type']})")
    print()

    # Check rules based on classification
    for rule in esg_rules:
        rule_id = rule.get('rule_id', '')
        severity = rule.get('severity', 'major').upper()

        # ESG_002: Engaging approach (no limits)
        if rule_id == 'ESG_002':
            if 'engaging' in classification_lower or 'engageante' in classification_lower:
                # No violations - unlimited ESG allowed
                pass

        # ESG_003: Reduced approach (< 10% volume)
        elif rule_id == 'ESG_003':
            if 'reduced' in classification_lower or 'rÃ©duite' in classification_lower:
                if esg_volume_percentage >= 10:
                    violations.append({
                        'type': 'ESG',
                        'severity': severity,
                        'slide': f"{len(esg_slides)} slides with ESG content",
                        'location': 'Document-wide',
                        'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                        'message': f"ESG content exceeds 10% limit for Reduced approach: {esg_volume_percentage:.1f}%",
                        'evidence': f"Total: {total_chars:,} chars, ESG: {esg_chars:,} chars. Must be < 10%. Primarily on slides: {', '.join([str(s['slide']) for s in esg_slides[:5]])}"
                    })

        # ESG_004: Prospectus-limited (NO ESG for retail)
        elif rule_id == 'ESG_004':
            if ('prospectus' in classification_lower or 'limitÃ©e' in classification_lower) and client_type.lower() == 'retail':
                if esg_slides:
                    esg_elements = []
                    for slide in esg_slides[:5]:
                        esg_elements.extend(slide.get('elements', []))

                    violations.append({
                        'type': 'ESG',
                        'severity': severity,
                        'slide': f"{len(esg_slides)} slides",
                        'location': ', '.join([f"Slide {s['slide']}" for s in esg_slides[:5]]),
                        'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                        'message': f"ESG mentions PROHIBITED for Prospectus-limited retail documents",
                        'evidence': f"Found ESG content on {len(esg_slides)} slides. Elements: {', '.join(set(esg_elements)[:8])}"
                    })

        # ESG_005: Other funds (only baseline exclusions)
        elif rule_id == 'ESG_005':
            if 'other' in classification_lower or 'autres' in classification_lower:
                # Check if content goes beyond baseline
                result = llm_check_esg_baseline(doc_text)

                if result.get('beyond_baseline') and result.get('confidence', 0) > 70:
                    violations.append({
                        'type': 'ESG',
                        'severity': severity,
                        'slide': 'Document-wide',
                        'location': 'ESG sections',
                        'rule': f"{rule_id}: {rule.get('rule_text', '')}",
                        'message': f"ESG content goes beyond baseline exclusions (prohibited for Other funds)",
                        'evidence': f"Topics beyond baseline: {', '.join(result.get('esg_topics_found', [])[:5])}. {result.get('reasoning', '')} Confidence: {result.get('confidence', 0)}%"
                    })

    return violations
def llm_compare_values_fuzzy(prospectus_value, document_value, value_type, field_name):
    """
    Use LLM to compare values with tolerance for formatting differences

    Args:
        prospectus_value: Value from prospectus
        document_value: Value from document
        value_type: 'percentage', 'currency', 'range', 'text', 'sri'
        field_name: Name of field being compared

    Returns:
        dict with matches, difference, confidence
    """
    if not tokenfactory_client:
        # Simple comparison fallback
        matches = str(prospectus_value).lower() in str(document_value).lower()
        return {'matches': matches, 'confidence': 70}

    prompt = f"""Compare these two values to determine if they are SUBSTANTIALLY the same.

FIELD: {field_name}
TYPE: {value_type}

PROSPECTUS VALUE: {prospectus_value}
DOCUMENT VALUE: {document_value}

COMPARISON RULES:
1. Allow formatting variations:
   - "0.75%" vs "0,75%" vs "0.75% per annum" vs "0.75% p.a."
   - "1,000,000" vs "1M" vs "1 million"
   - "4/7" vs "4 sur 7" vs "4 out of 7"

2. Allow language differences:
   - English vs French expressions
   - Different wording with same meaning

3. Allow minor rounding:
   - "0.75%" vs "0.7%" might be acceptable
   - "10-15%" vs "10% to 15%" is same

4. DO NOT allow substantive differences:
   - Different percentages (0.75% vs 1.5%)
   - Different ranges (10-20% vs 30-40%)
   - Missing information

Question: Are these values SUBSTANTIALLY the same?

Respond JSON:
{{
  "matches": true/false,
  "substantively_same": true/false,
  "difference": "description of any difference",
  "severity": "none" or "formatting_only" or "minor" or "major",
  "confidence": 0-100,
  "reasoning": "explanation"
}}"""

    try:
        response = tokenfactory_client.chat.completions.create(
            model="hosted_vllm/Llama-3.1-70B-Instruct",
            messages=[
                {"role": "system", "content": "You are a compliance analyst. Compare financial data values intelligently, allowing for formatting differences but flagging substantive discrepancies. Respond ONLY with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=500
        )

        result_text = response.choices[0].message.content.strip()
        result_text = result_text.replace('```json', '').replace('```', '').strip()
        result = json.loads(result_text)

        return result
    except Exception as e:
        print(f"    âš ï¸  LLM error comparing {field_name}: {e}")
        return {'matches': False, 'confidence': 50, 'error': str(e)}


def llm_check_prospectus_field(doc_text, field_name, expected_value, field_type='text'):
    """
    Check if expected prospectus field is present in document

    Args:
        doc_text: Document text
        field_name: Name of field (e.g., 'SRI', 'management_fees')
        expected_value: Expected value from prospectus
        field_type: Type of field

    Returns:
        dict with found, found_value, matches, confidence
    """
    if not tokenfactory_client:
        # Simple search fallback
        found = str(expected_value).lower() in doc_text.lower()
        return {'found': found, 'found_value': expected_value if found else None, 'matches': found, 'confidence': 60}

    prompt = f"""Find and compare this prospectus field in the document.

FIELD: {field_name}
EXPECTED VALUE (from prospectus): {expected_value}
TYPE: {field_type}

DOCUMENT TEXT (excerpt):
{doc_text[:2500]}

Tasks:
1. Search for this field in the document (allow for different wording/languages)
2. Extract the actual value shown
3. Compare if it matches the expected value (allow formatting differences)

Example field mappings:
- "management_fees" â†’ look for "frais de gestion", "management fee", "fees", "charges"
- "SRI" â†’ look for "SRI", "SRRI", "indicateur de risque", "risk indicator"
- "minimum_investment" â†’ look for "minimum", "souscription minimale", "minimum subscription"

Respond JSON:
{{
  "found": true/false,
  "found_value": "actual value in document or null",
  "found_where": "description of where found",
  "matches_expected": true/false,
  "confidence": 0-100,
  "reasoning": "explanation"
}}"""

    try:
        response = tokenfactory_client.chat.completions.create(
            model="hosted_vllm/Llama-3.1-70B-Instruct",
            messages=[
                {"role": "system", "content": "You are a multilingual document analyst. Find and compare financial data fields. Understand semantic equivalents. Respond ONLY with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=600
        )

        result_text = response.choices[0].message.content.strip()
        result_text = result_text.replace('```json', '').replace('```', '').strip()
        result = json.loads(result_text)

        return result
    except Exception as e:
        print(f"    âš ï¸  LLM error checking {field_name}: {e}")
        return {'found': False, 'confidence': 50}


def check_prospectus_compliance_enhanced(doc, prospectus_data):
    """
    Enhanced prospectus compliance checker with LLM fuzzy matching
    """
    violations = []

    if not prospectus_data or prospectus_data == {'raw_text': prospectus_data.get('raw_text', '')}:
        return violations

    doc_text = extract_all_text_from_doc(doc)

    print("\nðŸ” Checking prospectus compliance with LLM fuzzy matching...")

    # Track AI confidence for manual review decision
    low_confidence_checks = []
    contradictions = []

    # PROSP_003: SRI Check
    if 'sri' in prospectus_data and prospectus_data['sri']:
        prospectus_sri = prospectus_data['sri']

        result = llm_check_prospectus_field(doc_text, 'SRI', prospectus_sri, 'sri')

        if not result.get('found') or result.get('confidence', 0) < 70:
            violations.append({
                'type': 'PROSPECTUS',
                'severity': 'CRITICAL',
                'slide': 'SRI section',
                'location': 'Missing',
                'rule': 'PROSP_003: SRI must match prospectus',
                'message': f"SRI not found or unclear in document",
                'evidence': f"Prospectus SRI: {prospectus_sri}. {result.get('reasoning', '')} Confidence: {result.get('confidence', 0)}%"
            })
        elif not result.get('matches_expected'):
            # Found but doesn't match - use fuzzy comparison
            comparison = llm_compare_values_fuzzy(
                prospectus_sri,
                result.get('found_value'),
                'sri',
                'SRI'
            )

            if not comparison.get('substantively_same') or comparison.get('severity') in ['major']:
                violations.append({
                    'type': 'PROSPECTUS',
                    'severity': 'CRITICAL',
                    'slide': 'SRI section',
                    'location': result.get('found_where', 'Document'),
                    'rule': 'PROSP_003: SRI must match prospectus exactly',
                    'message': f"SRI mismatch: Document '{result.get('found_value')}' â‰  Prospectus '{prospectus_sri}'",
                    'evidence': f"{comparison.get('difference', 'Values differ')}. Confidence: {comparison.get('confidence', 0)}%"
                })

    # PROSP_013: Management Fees
    if 'management_fees' in prospectus_data and prospectus_data['management_fees']:
        prospectus_fees = prospectus_data['management_fees']

        result = llm_check_prospectus_field(doc_text, 'management_fees', prospectus_fees, 'percentage')

        if result.get('found') and result.get('found_value'):
            # Compare values
            comparison = llm_compare_values_fuzzy(
                prospectus_fees,
                result.get('found_value'),
                'percentage',
                'Management Fees'
            )

            if not comparison.get('substantively_same'):
                severity = 'CRITICAL' if comparison.get('severity') == 'major' else 'MAJOR'
                violations.append({
                    'type': 'PROSPECTUS',
                    'severity': severity,
                    'slide': 'Fees section',
                    'location': result.get('found_where', 'Document'),
                    'rule': 'PROSP_013: Management fees must match prospectus',
                    'message': f"Fees mismatch: Document '{result.get('found_value')}' vs Prospectus '{prospectus_fees}'",
                    'evidence': f"{comparison.get('difference', '')}. {comparison.get('reasoning', '')} Confidence: {comparison.get('confidence', 0)}%"
                })
        elif 'fees' in doc_text.lower() or 'frais' in doc_text.lower():
            violations.append({
                'type': 'PROSPECTUS',
                'severity': 'MAJOR',
                'slide': 'Fees section',
                'location': 'Unclear',
                'rule': 'PROSP_013: Management fees must match prospectus',
                'message': f"Fees mentioned but value unclear or doesn't match prospectus",
                'evidence': f"Expected: {prospectus_fees}. Confidence: {result.get('confidence', 0)}%"
            })

    # PROSP_012: Minimum Investment
    if 'minimum_investment' in prospectus_data and prospectus_data['minimum_investment']:
        prospectus_min = prospectus_data['minimum_investment']

        result = llm_check_prospectus_field(doc_text, 'minimum_investment', prospectus_min, 'currency')

        if result.get('found') and result.get('found_value'):
            comparison = llm_compare_values_fuzzy(
                prospectus_min,
                result.get('found_value'),
                'currency',
                'Minimum Investment'
            )

            if not comparison.get('substantively_same'):
                violations.append({
                    'type': 'PROSPECTUS',
                    'severity': 'MAJOR',
                    'slide': 'Fund characteristics',
                    'location': result.get('found_where'),
                    'rule': 'PROSP_012: Minimum investment must match prospectus',
                    'message': f"Minimum investment mismatch: '{result.get('found_value')}' vs '{prospectus_min}'",
                    'evidence': f"{comparison.get('difference', '')}. Confidence: {comparison.get('confidence', 0)}%"
                })

    # PROSP_004: Benchmark
    if 'benchmark' in prospectus_data and prospectus_data['benchmark']:
        prospectus_benchmark = prospectus_data['benchmark']

        # Check if performance is shown
        has_performance = any(kw in doc_text.lower() for kw in ['performance', 'rendement', 'return'])

        if has_performance:
            result = llm_check_prospectus_field(doc_text, 'benchmark', prospectus_benchmark, 'text')

            if not result.get('found') or result.get('confidence', 0) < 70:
                violations.append({
                    'type': 'PROSPECTUS',
                    'severity': 'CRITICAL',
                    'slide': 'Performance section',
                    'location': 'Missing benchmark',
                    'rule': 'PROSP_004: Must use official prospectus benchmark',
                    'message': 'Performance shown without prospectus benchmark',
                    'evidence': f"Required benchmark: {prospectus_benchmark[:100]}. Confidence: {result.get('confidence', 0)}%"
                })

    # Track low confidence checks for potential manual review
    for field_name, field_value in prospectus_data.items():
        if field_name != 'raw_text' and field_value:
            result = llm_check_prospectus_field(doc_text, field_name, field_value)
            if result.get('confidence', 100) < 80:
                low_confidence_checks.append({
                    'field': field_name,
                    'confidence': result.get('confidence', 0),
                    'found': result.get('found'),
                    'reasoning': result.get('reasoning', '')
                })

    # Conditional PROSP_008: Only flag if confidence is genuinely low or contradictions found
    if len(low_confidence_checks) >= 3 or contradictions:
        evidence_parts = []
        if low_confidence_checks:
            evidence_parts.append(f"Low confidence on {len(low_confidence_checks)} fields: " +
                                ', '.join([f"{c['field']} ({c['confidence']}%)" for c in low_confidence_checks[:3]]))
        if contradictions:
            evidence_parts.append(f"Contradictions detected: {', '.join(contradictions)}")

        violations.append({
            'type': 'PROSPECTUS',
            'severity': 'WARNING',
            'slide': 'Document-wide',
            'location': 'Multiple data points',
            'rule': 'PROSP_008: Manual verification recommended',
            'message': 'âš ï¸  Some prospectus comparisons have low confidence - manual review recommended',
            'evidence': ' | '.join(evidence_parts)
        })

    print(f"âœ“ Prospectus compliance check complete: {len(violations)} violations, {len(low_confidence_checks)} low-confidence checks\n")

    return violations
# ============================================================================
# ENHANCED MAIN COMPLIANCE CHECKER - Replace entire check_document_compliance
# ============================================================================

def check_document_compliance(json_file_path):
    """
    Enhanced document compliance checker with LLM-powered context-aware validation

    All checks now use Token Factory Llama 3.1 for:
    - Context-aware country detection
    - Multilingual text understanding
    - Fuzzy value matching
    - Semantic phrase detection
    - Multi-modal performance detection

    Returns:
        dict with violations list and confidence scores
    """
    try:
        # Load document
        with open(json_file_path, 'r', encoding='utf-8') as f:
            doc = json.load(f)

        violations = []

        # Extract metadata
        doc_metadata = doc.get('document_metadata', {})
        fund_isin = doc_metadata.get('fund_isin')
        client_type = doc_metadata.get('client_type', 'retail')
        doc_type = doc_metadata.get('document_type', 'fund_presentation')
        fund_status = doc_metadata.get('fund_status', 'active')
        esg_classification = doc_metadata.get('fund_esg_classification', 'other')
        country_code = doc_metadata.get('country_code', None)
        fund_age_years = doc_metadata.get('fund_age_years', None)

        print(f"\n{'='*70}")
        print(f"ðŸ” ENHANCED COMPLIANCE REPORT (LLM-Powered)")
        print(f"{'='*70}")
        print(f"Fund ISIN: {fund_isin}")
        print(f"Client Type: {client_type.upper()}")
        print(f"Document Type: {doc_type}")
        print(f"Fund Status: {fund_status}")
        print(f"ESG Classification: {esg_classification}")
        if tokenfactory_client:
            print(f"âœ“ LLM: Token Factory (Llama-3.1-70B)")
        else:
            print(f"âš ï¸  LLM: Not available (fallback to basic matching)")
        print(f"{'='*70}\n")

        # ====================================================================
        # CHECK 1: REGISTRATION (Enhanced with LLM Context Analysis)
        # ====================================================================
        if fund_isin and funds_db:
            fund_info = None
            for fund in funds_db:
                if fund['isin'] == fund_isin:
                    fund_info = fund
                    break

            if fund_info:
                authorized_countries = fund_info['authorized_countries']
                print(f"âœ“ Fund: {fund_info['fund_name']}")
                print(f"âœ“ Authorized in {len(authorized_countries)} countries\n")

                # Use enhanced LLM-based registration check
                reg_violations = check_registration_rules_enhanced(
                    doc,
                    fund_isin,
                    authorized_countries
                )
                violations.extend(reg_violations)

                if not reg_violations:
                    print("âœ… Registration compliance: OK\n")
            else:
                print(f"âš ï¸  Fund {fund_isin} not found in registration database\n")

        # ====================================================================
        # CHECK 2: DISCLAIMERS (Keep existing AI check - it's already good)
        # ====================================================================
        if disclaimers_db:
            doc_type_mapping = {
                'fund_presentation': 'OBAM Presentation',
                'commercial_doc': 'Commercial documentation',
                'fact_sheet': 'OBAM Presentation'
            }

            disclaimer_type = doc_type_mapping.get(doc_type, 'OBAM Presentation')

            if disclaimer_type in disclaimers_db:
                client_key = 'professional' if client_type.lower() == 'professional' else 'retail'
                required_disclaimer = disclaimers_db[disclaimer_type].get(client_key)

                if required_disclaimer and len(required_disclaimer) > 50:
                    all_slides_text = extract_all_text_from_doc(doc)

                    if model:  # Using Gemini for this check
                        result = check_disclaimer_in_document(all_slides_text, required_disclaimer)

                        if result.get('status') == 'MISSING':
                            violations.append({
                                'type': 'DISCLAIMER',
                                'severity': 'CRITICAL',
                                'slide': 'Document-wide',
                                'location': 'Missing',
                                'rule': f'Required {client_key} disclaimer',
                                'message': f"Required disclaimer not found",
                                'evidence': result.get('explanation', ''),
                                'confidence': result.get('confidence', 0)
                            })

        # ====================================================================
        # CHECK 3: STRUCTURE (Enhanced with LLM)
        # ====================================================================
        if structure_rules:
            print("ðŸ” Checking structure with LLM...\n")
            structure_violations = check_structure_rules_enhanced(
                doc,
                client_type,
                fund_status,
                prospectus_data
            )
            violations.extend(structure_violations)

            if not structure_violations:
                print("âœ… Structure compliance: OK\n")

        # ====================================================================
        # CHECK 4: SECURITIES/VALUES (Enhanced with LLM)
        # ====================================================================
        if values_rules:
            print("ðŸ” Checking securities/values with LLM context analysis...\n")
            values_violations = check_values_rules_enhanced(doc)
            violations.extend(values_violations)

            if not values_violations:
                print("âœ… Securities/Values compliance: OK\n")

        # ====================================================================
        # CHECK 5: ESG (Enhanced with LLM Section Analysis)
        # ====================================================================
        if esg_rules:
            esg_violations = check_esg_rules_enhanced(doc, esg_classification, client_type)
            violations.extend(esg_violations)

            if not esg_violations:
                print("âœ… ESG compliance: OK\n")

        # ====================================================================
        # CHECK 6: PERFORMANCE (Enhanced with LLM)
        # ====================================================================
        if performance_rules:
            perf_violations = check_performance_rules_enhanced(
                doc,
                client_type,
                fund_age_years
            )
            violations.extend(perf_violations)

            if not perf_violations:
                print("âœ… Performance compliance: OK\n")

        # ====================================================================
        # CHECK 7: GENERAL RULES (Enhanced with LLM)
        # ====================================================================
        if general_rules:
            print("ðŸ” Checking general rules with LLM...\n")
            gen_violations = check_general_rules_enhanced(doc, client_type, country_code)
            violations.extend(gen_violations)

            if not gen_violations:
                print("âœ… General rules compliance: OK\n")

        # ====================================================================
        # CHECK 8: PROSPECTUS (Enhanced with LLM Fuzzy Matching)
        # ====================================================================
        if prospectus_data and prospectus_rules:
            prosp_violations = check_prospectus_compliance_enhanced(doc, prospectus_data)
            violations.extend(prosp_violations)

            if not prosp_violations:
                print("âœ… Prospectus compliance: OK\n")

        # ====================================================================
        # FINAL REPORT
        # ====================================================================
        print(f"\n{'='*70}")
        if len(violations) == 0:
            print("âœ… NO VIOLATIONS FOUND - Document is compliant!")
        else:
            print(f"âŒ {len(violations)} VIOLATION(S) FOUND")
        print(f"{'='*70}\n")

        # Display violations with confidence scores
        for i, v in enumerate(violations, 1):
            print(f"{'â”'*70}")
            print(f"[{v['severity']}] {v['type']} Violation #{i} - {v['slide']}")
            print(f"{'â”'*70}")
            print(f"ðŸ“‹ Rule: {v['rule']}")
            print(f"âš ï¸  Issue: {v['message']}")
            print(f"ðŸ“ Location: {v['location']}")

            # Show confidence if available
            if 'confidence' in v:
                conf_emoji = 'ðŸŽ¯' if v['confidence'] >= 80 else 'âš¡' if v['confidence'] >= 60 else 'âš ï¸'
                print(f"{conf_emoji} Confidence: {v['confidence']}%")

            print(f"\nðŸ“„ Evidence:")
            print(f"   {v['evidence']}")
            print()

        # Summary statistics
        if violations:
            print(f"\n{'='*70}")
            print(f"SUMMARY")
            print(f"{'='*70}")

            # Count by type
            type_counts = {}
            for v in violations:
                vtype = v['type']
                type_counts[vtype] = type_counts.get(vtype, 0) + 1

            print(f"\nViolations by type:")
            for vtype, count in sorted(type_counts.items()):
                print(f"  â€¢ {vtype}: {count}")

            # Count by severity
            severity_counts = {}
            for v in violations:
                sev = v['severity']
                severity_counts[sev] = severity_counts.get(sev, 0) + 1

            print(f"\nViolations by severity:")
            for sev in ['CRITICAL', 'MAJOR', 'WARNING']:
                if sev in severity_counts:
                    print(f"  â€¢ {sev}: {severity_counts[sev]}")

            # Average confidence
            confidences = [v.get('confidence', 0) for v in violations if 'confidence' in v]
            if confidences:
                avg_conf = sum(confidences) / len(confidences)
                print(f"\nAverage confidence: {avg_conf:.1f}%")

        return {
            'total_violations': len(violations),
            'violations': violations,
            'fund_isin': fund_isin,
            'client_type': client_type,
            'llm_enabled': tokenfactory_client is not None
        }

    except Exception as e:
        print(f"\nâŒ Error checking document: {e}")
        import traceback
        traceback.print_exc()
        return {'error': str(e)}


# ============================================================================
# STEP 13: Interactive Compliance Chatbot
# ============================================================================

def compliance_chatbot():
    """AI-powered interactive compliance checker"""
    print("\n" + "="*70)
    print("AI COMPLIANCE CHECKER - INTERACTIVE MODE")
    print("="*70)
    print("\nðŸ” Available checks:")
    print("  1. Full document compliance (JSON file)")
    print("     - Registration (unauthorized countries)")
    print("     - Disclaimers (required text)")
    print("     - Structure (cover page, slide 2, back page)")
    print("     - Securities/Values (investment recommendations)")
    print("  2. Fund registration lookup")
    print("  3. Disclaimer reference lookup")
    print("  4. Structure rules reference")
    print("  5. Securities/Values rules reference")
    print("  6. Performance rules reference")
    print("  7. General rules reference")
    print("  8. Prospectus rules reference")

    print("\nðŸ“ Commands:")
    print("  check json [filename.json]        - Full compliance check")
    print("  check [fund] in [countries]       - Check registration only")
    print("  show disclaimers                  - Show required disclaimers")
    print("  show structure rules              - Show structure requirements")
    print("  show performance rules            - Show performance requirements")
    print("  show general rules                - Show general requirements")
    print("  show prospectus rules            - Show prospectus requirements")
    print("  show prospectus                   - Show parsed prospectus data")
    print("  show values rules                 - Show securities/values rules")
    print("  search [fund]                     - Find fund details")
    print("  stats                             - Database statistics")
    print("  quit                              - Exit")
    print("\nðŸ’¡ Examples:")
    print("  check json exemple_performance.json")
    print("  check ODDO BHF Alpha Europe in Germany, United States")
    print("  show values rules\n")

    while True:
        user_input = input("\nðŸ’¬ You: ").strip()

        if not user_input:
            continue

        if user_input.lower() == 'quit':
            print("ðŸ‘‹ Goodbye!")
            break

        # ================================================================
        # CHECK JSON DOCUMENT COMPLIANCE
        # ================================================================
        if user_input.lower().startswith('check json'):
            json_filename = user_input.replace('check json', '').strip()

            if not json_filename:
                print("âŒ Please specify JSON filename")
                print("   Example: check json exemple_performance.json")
                continue

            if not os.path.exists(json_filename):
                print(f"âŒ File '{json_filename}' not found in current directory")
                print("ðŸ“ Uploading file...")
                uploaded = files.upload()
                if uploaded:
                    json_filename = list(uploaded.keys())[0]
                else:
                    print("âŒ No file uploaded")
                    continue

            result = check_document_compliance(json_filename)

            if 'error' not in result:
                print(f"\n{'='*70}")
                print("ðŸ“Š COMPLIANCE CHECK COMPLETE")
                print(f"{'='*70}")
                print(f"Total violations: {result['total_violations']}")

                if result['total_violations'] == 0:
                    print("\nðŸŽ‰ Document is fully compliant!")
                    print("   âœ… Registration: OK")
                    print("   âœ… Disclaimers: OK")
                    print("   âœ… Structure: OK")
                    print("   âœ… Securities/Values: OK")
                else:
                    print(f"\nâš ï¸  Please review and fix {result['total_violations']} violation(s) above")

            continue

        # ================================================================
        # SHOW VALUES RULES
        # ================================================================
        if user_input.lower() == 'show values rules':
            if not values_rules:
                print("âŒ Values rules not loaded")
                continue

            print(f"\nðŸ“‹ SECURITIES/VALUES RULES ({len(values_rules)} total)\n")
            print("These rules prevent investment recommendations (MAR regulation)\n")

            # Show only prohibition rules (not "allowed" ones)
            prohibition_rules = [r for r in values_rules if r.get('category') != 'allowed']
            allowed_rules = [r for r in values_rules if r.get('category') == 'allowed']

            print(f"{'='*70}")
            print("ðŸš« PROHIBITIONS (What you CANNOT do)")
            print(f"{'='*70}")

            for rule in prohibition_rules[:10]:  # Show first 10
                print(f"\n  {rule.get('rule_id', '')} [{rule.get('severity', '').upper()}]")
                print(f"  âž¤ {rule.get('rule_text', '')}")

                prohibited = rule.get('prohibited_phrases', [])
                if prohibited:
                    print(f"  âŒ Prohibited phrases: {', '.join(prohibited[:5])}{'...' if len(prohibited) > 5 else ''}")

            print(f"\n{'='*70}")
            print("âœ… ALLOWED CONTENT")
            print(f"{'='*70}")

            for rule in allowed_rules[:5]:
                print(f"\n  {rule.get('rule_id', '')} - {rule.get('rule_text', '')}")

            print(f"\n{'='*70}\n")
            print("ðŸ’¡ Key principle: NO buy/sell recommendations, valuations, or projections")
            continue

        # ================================================================
        # SHOW STRUCTURE RULES
        # ================================================================
        if user_input.lower() == 'show structure rules':
            if not structure_rules:
                print("âŒ Structure rules not loaded")
                continue

            print(f"\nðŸ“‹ STRUCTURE RULES ({len(structure_rules)} total)\n")

            sections = {}
            for rule in structure_rules:
                section = rule.get('section', 'Other')
                if section not in sections:
                    sections[section] = []
                sections[section].append(rule)

            for section, rules in sections.items():
                print(f"\n{'='*70}")
                print(f"ðŸ“ {section.upper()}")
                print(f"{'='*70}")

                for rule in rules:
                    print(f"\n  {rule.get('rule_id', '')} [{rule.get('severity', '').upper()}]")
                    print(f"  âž¤ {rule.get('rule_text', '')}")
                    print(f"  ðŸ“ {rule.get('description', '')}")

                    applies_to = rule.get('applies_to', {})
                    print(f"  ðŸŽ¯ Applies to: {', '.join(applies_to.get('client_type', []))}")

                    if applies_to.get('fund_status', []) != ['all']:
                        print(f"  âš ï¸  Special status: {', '.join(applies_to.get('fund_status', []))}")

            print(f"\n{'='*70}\n")
            continue
        # ================================================================
        # SHOW ESG RULES
        # ================================================================
        if user_input.lower() == 'show esg rules':
            if not esg_rules:
                print("âŒ ESG rules not loaded")
                continue

            print(f"\nðŸ“‹ ESG RULES ({len(esg_rules)} total, ESG_001 excluded)\n")

            for rule in esg_rules:
                print(f"\n{'='*70}")
                print(f"{rule.get('rule_id', '')} [{rule.get('severity', '').upper()}]")
                print(f"{'='*70}")
                print(f"âž¤ {rule.get('rule_text', '')}")
                print(f"ðŸ“ {rule.get('detailed_description', '')}")

            print(f"\n{'='*70}\n")
            continue
          # ================================================================
        # SHOW PERFORMANCE RULES
        # ================================================================
        if user_input.lower() == 'show performance rules':
            if not performance_rules:
                print("âŒ Performance rules not loaded")
                continue

            print(f"\nðŸ“‹ PERFORMANCE RULES ({len(performance_rules)} total)\n")

            # Group by category
            categories = {}
            for rule in performance_rules:
                cat = rule.get('category', 'other')
                if cat not in categories:
                    categories[cat] = []
                categories[cat].append(rule)

            for cat_name, rules in categories.items():
                print(f"\n{'='*70}")
                print(f"ðŸ“ {cat_name.upper().replace('_', ' ')}")
                print(f"{'='*70}")

                for rule in rules[:5]:  # Show first 5 per category
                    print(f"\n  {rule.get('rule_id', '')} [{rule.get('severity', '').upper()}]")
                    print(f"  âž¤ {rule.get('rule_text', '')[:150]}...")

                if len(rules) > 5:
                    print(f"\n  ... and {len(rules) - 5} more rules")

            print(f"\n{'='*70}\n")
            print("ðŸ’¡ Tip: Type 'show perf [rule_id]' for details (e.g., 'show perf PERF_001')")
            continue
        # ================================================================
        # SHOW DISCLAIMERS
        # ================================================================
        if user_input.lower() == 'show disclaimers':
            if not disclaimers_db:
                print("âŒ Disclaimers database not loaded")
                continue

            print(f"\nðŸ“‹ Available disclaimer types ({len(disclaimers_db)}):\n")
            for i, doc_type in enumerate(disclaimers_db.keys(), 1):
                print(f"  {i}. {doc_type}")

            print("\nðŸ’¡ Tip: Type the document type name to see the full disclaimer text")
            continue
        # ================================================================
        # SHOW PROSPECTUS DATA
        # ================================================================
        if user_input.lower() == 'show prospectus':
            if not prospectus_data:
                print("âŒ Prospectus not loaded")
                continue

            print(f"\nðŸ“„ PROSPECTUS DATA\n")
            print(f"{'='*70}")
            for key, value in prospectus_data.items():
                if key != 'raw_text':
                    print(f"{key}: {value}")
            print(f"{'='*70}\n")
            continue
        # ================================================================
        # SHOW GENERAL RULES
        # ================================================================
        if user_input.lower() == 'show general rules':
            if not general_rules:
                print("âŒ General rules not loaded")
                continue

            print(f"\nðŸ“‹ GENERAL RULES ({len(general_rules)} total)\n")

            # Group by section
            sections = {}
            for rule in general_rules:
                sec = rule.get('section', '1 - General Rules')
                if sec not in sections:
                    sections[sec] = []
                sections[sec].append(rule)

            for sec_name, rules in sections.items():
                print(f"\n{'='*70}")
                print(f"ðŸ“ {sec_name.upper()}")
                print(f"{'='*70}")

                for rule in rules[:8]:  # Show first 8 per section
                    print(f"\n  {rule.get('rule_id', '')} [{rule.get('severity', '').upper()}]")
                    print(f"  âž¤ {rule.get('rule_text', '')}")

                if len(rules) > 8:
                    print(f"\n  ... and {len(rules) - 8} more rules")

            print(f"\n{'='*70}\n")
            continue
        # ================================================================
        # SHOW SPECIFIC DISCLAIMER
        # ================================================================
        if any(dt.lower() in user_input.lower() for dt in disclaimers_db.keys()):
            for doc_type in disclaimers_db.keys():
                if doc_type.lower() in user_input.lower():
                    is_prof = metadata.get('Le client est-il un professionnel', False) if metadata else False
                    client_type = 'professional' if is_prof else 'retail'

                    disclaimer = disclaimers_db[doc_type].get(client_type)

                    print(f"\nðŸ“„ Document type: {doc_type}")
                    print(f"ðŸ‘¤ Client type: {client_type.upper()}")
                    print(f"\n{'='*70}")
                    print("REQUIRED DISCLAIMER TEXT:")
                    print(f"{'='*70}\n")
                    print(disclaimer if disclaimer else "âŒ No disclaimer defined for this client type")
                    print(f"\n{'='*70}")
                    break
            continue

        # ================================================================
        # CHECK REGISTRATION
        # ================================================================
        if user_input.lower().startswith('check') and ' in ' in user_input:
            parts = user_input.split(' in ')
            if len(parts) != 2:
                print("âŒ Format: check [fund_name/ISIN] in [country1, country2]")
                continue

            fund_name = parts[0].replace('check', '').strip()
            countries = [c.strip() for c in parts[1].split(',')]

            matches = []
            search_term = fund_name.upper()

            for fund in funds_db:
                if (search_term in fund['fund_name'].upper() or
                    search_term in fund.get('isin', '').upper() or
                    search_term in fund.get('share_class', '').upper()):
                    matches.append(fund)

            if not matches:
                print(f"âŒ Fund '{fund_name}' not found in database")
                continue

            for fund in matches[:3]:
                print(f"\nðŸ“Š Fund: {fund['fund_name']}")
                print(f"   Family: {fund['fund_family']}")
                print(f"   ISIN: {fund['isin']}")
                print(f"   Share class: {fund['share_class']}")

                violations = []
                authorized = []

                for country in countries:
                    found = False
                    for auth_country in fund['authorized_countries']:
                        country_clean = country.lower().replace(' (fund)', '').strip()
                        auth_clean = auth_country.lower().replace(' (fund)', '').strip()

                        if country_clean == auth_clean or country_clean in auth_clean or auth_clean in country_clean:
                            authorized.append(country)
                            found = True
                            break

                    if not found:
                        violations.append(country)

                if violations:
                    print(f"\n   âŒ NOT AUTHORIZED in: {', '.join(violations)}")
                else:
                    print(f"\n   âœ… AUTHORIZED in all requested countries")

                if authorized:
                    print(f"   âœ“ Authorized: {', '.join(authorized)}")

                print(f"\n   Full authorization list ({len(fund['authorized_countries'])} countries):")
                print(f"   {', '.join(sorted(fund['authorized_countries']))}")

            continue

        # ================================================================
        # SEARCH FUND
        # ================================================================
        if user_input.lower().startswith('search'):
            if not funds_db:
                print("âŒ Registration database not loaded")
                continue

            search_term = user_input.replace('search', '').strip().upper()
            matches = [f for f in funds_db if search_term in f['fund_name'].upper() or search_term in f.get('isin', '').upper()]

            if matches:
                print(f"\nâœ“ Found {len(matches)} matching fund(s):\n")
                for fund in matches[:10]:
                    print(f"  â€¢ {fund['fund_name']} ({fund['share_class']})")
                    print(f"    Family: {fund['fund_family']}")
                    print(f"    ISIN: {fund['isin']}")
                    print(f"    Authorized in {len(fund['authorized_countries'])} countries\n")
                if len(matches) > 10:
                    print(f"  ... and {len(matches) - 10} more")
            else:
                print(f"âŒ No funds found matching '{search_term}'")
            continue


        # ================================================================
        # STATS
        # ================================================================
        if user_input.lower() == 'stats':
            print(f"\nðŸ“Š DATABASE STATISTICS")
            print(f"{'='*70}")

            if funds_db:
                print(f"\n  Registration Database:")
                print(f"  â€¢ Total fund share classes: {len(funds_db)}")
                print(f"  â€¢ Countries tracked: {len(all_countries)}")

                families = {}
                for fund in funds_db:
                    families[fund['fund_family']] = families.get(fund['fund_family'], 0) + 1

                print(f"  â€¢ Fund families: {len(families)}")
                print(f"\n  Top fund families:")
                for family, count in sorted(families.items(), key=lambda x: x[1], reverse=True)[:5]:
                    print(f"    - {family}: {count} share classes")

            if disclaimers_db:
                print(f"\n  Disclaimers Database:")
                print(f"  â€¢ Document types: {len(disclaimers_db)}")
                print(f"  â€¢ Types available: {', '.join(list(disclaimers_db.keys())[:3])}...")

            if structure_rules:
                print(f"\n  Structure Rules:")
                print(f"  â€¢ Total rules: {len(structure_rules)}")

                severities = {}
                for rule in structure_rules:
                    sev = rule.get('severity', 'unknown')
                    severities[sev] = severities.get(sev, 0) + 1

                print(f"  â€¢ By severity:")
                for sev, count in sorted(severities.items()):
                    print(f"    - {sev.upper()}: {count} rules")

                sections = {}
                for rule in structure_rules:
                    sec = rule.get('section', 'unknown')
                    sections[sec] = sections.get(sec, 0) + 1

                print(f"  â€¢ By section:")
                for sec, count in sections.items():
                    print(f"    - {sec}: {count} rules")

            if values_rules:
                print(f"\n  Securities/Values Rules:")
                print(f"  â€¢ Total rules: {len(values_rules)}")

                categories = {}
                for rule in values_rules:
                    cat = rule.get('category', 'unknown')
                    categories[cat] = categories.get(cat, 0) + 1

                print(f"  â€¢ By category:")
                for cat, count in categories.items():
                    print(f"    - {cat}: {count} rules")
            if esg_rules:
                print(f"\n  ESG Rules:")
                print(f"  â€¢ Total rules: {len(esg_rules)} (ESG_001 excluded)")

            if performance_rules:
                print(f"\n  Performance Rules:")
                print(f"  â€¢ Total rules: {len(performance_rules)}")

                categories = {}
                for rule in performance_rules:
                    cat = rule.get('category', 'unknown')
                    categories[cat] = categories.get(cat, 0) + 1

                print(f"  â€¢ By category:")
                for cat, count in sorted(categories.items()):
                    print(f"    - {cat}: {count} rules")
            if general_rules:
                print(f"\n  General Rules:")
                print(f"  â€¢ Total rules: {len(general_rules)}")

                severities = {}
                for rule in general_rules:
                    sev = rule.get('severity', 'unknown')
                    severities[sev] = severities.get(sev, 0) + 1

                print(f"  â€¢ By severity:")
                for sev, count in sorted(severities.items()):
                    print(f"    - {sev}: {count} rules")
            if metadata:
                print(f"\n  Current Metadata:")
                print(f"  â€¢ Client: {'Professional' if metadata.get('Le client est-il un professionnel') else 'Retail'}")
                print(f"  â€¢ New Product: {'Yes' if metadata.get('Le document fait-il rÃ©fÃ©rence Ã  un nouveau Produit') else 'No'}")
                print(f"  â€¢ Management Company: {metadata.get('SociÃ©tÃ© de Gestion', 'N/A')}")

            print(f"\n{'='*70}")
            continue

        # ================================================================
        # UNKNOWN COMMAND
        # ================================================================
        print("â“ Unknown command. Available commands:")
        print("   check json [file]  |  check [fund] in [countries]")
        print("   show disclaimers   |  show structure rules  |  show values rules")
        print("   search [fund]  |  stats  |  quit")

# ============================================================================
# STEP 14: Start the Chatbot
# ============================================================================